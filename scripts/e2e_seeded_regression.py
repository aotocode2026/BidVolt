"""种子态回归 E2E：补齐历史账号/输入矩阵/错误路径三缺口（Issue #13 复盘产物）。

用法：
    .venv\\Scripts\\python scripts/e2e_seeded_regression.py --base http://47.100.182.3:28123 --tag seeded

场景（与常规 e2e_browser_demo.py 的"全新账号 26 步"互补）：
1. 种子历史账号：注册 → 建 3 个项目 → 每个项目 evaluate 产生评分记录 →
   项目列表必须 200 且各项目取到最新评分（此前 PostgreSQL 基数约束 500 的精确回归）；
2. 浏览器登录该历史账号（模拟产品换浏览器重登）→ 项目列表 UI 无"项目列表失败"、3 行可见；
3. 多格式输入矩阵：txt/docx/xlsx/pptx/ofd(/pdf) 全部解析成功（status=3），
   损坏 docx（合法 zip 容器+垃圾内容）解析失败且上传时即显示具体原因；
4. 错误路径：含损坏文件的项目 UI 触发招标解析 → 任务失败日志必须含具体原因、
   禁止"未知错误"；API error.message 与 SSE failed 事件同样携带原因。
"""
from __future__ import annotations

import argparse
import io
import re
import sys
import time
import zipfile
from pathlib import Path

import httpx
from playwright.sync_api import sync_playwright

BASE = "http://127.0.0.1:8123"

results: list[tuple[str, bool, str]] = []


def record(name: str, ok: bool, note: str = "") -> None:
    results.append((name, ok, note))
    print(f"[{'PASS' if ok else 'FAIL'}] {name} {('(' + note + ')') if note else ''}", flush=True)


def wait_log(page, text: str, timeout_ms: int = 60000) -> bool:
    try:
        page.wait_for_function(
            "t => document.getElementById('log').innerText.includes(t)",
            arg=text,
            timeout=timeout_ms,
        )
        return True
    except Exception:  # noqa: BLE001
        return False


def log_text(page) -> str:
    return page.evaluate("document.getElementById('log').innerText")


def tab(page, label: str) -> None:
    page.click(f"#tabs button:has-text('{label}')")
    page.wait_for_timeout(300)


def api(base: str, method: str, path: str, token: str | None = None, json=None, files=None, data=None, timeout=60):
    headers = {}
    if token:
        headers["Authorization"] = f"Bearer {token}"
    r = httpx.request(method, f"{base}{path}", headers=headers, json=json, files=files, data=data, timeout=timeout)
    r.raise_for_status()
    return r.json()


def poll_task(base: str, token: str, task_id: int, timeout_s: int) -> dict:
    deadline = time.monotonic() + timeout_s
    while time.monotonic() < deadline:
        try:
            r = httpx.get(f"{base}/api/v1/tasks/{task_id}", headers={"Authorization": f"Bearer {token}"}, timeout=10)
            if r.status_code == 200 and r.json().get("status") in (3, 6):
                return r.json()
        except Exception:  # noqa: BLE001
            pass
        time.sleep(3)
    return {"status": -1}


# ---------- 输入矩阵文件构造 ----------

def make_txt() -> bytes:
    return (
        "招标公告\n一、资质要求：投标人须具备电力工程施工总承包资质。\n"
        "二、技术要求：电缆 YJV-3x95 需符合 GB/T 12706 标准。\n"
        "三、商务要求：投标保证金 2 万元，履约保证金 5%。\n"
    ).encode()


def make_docx() -> bytes:
    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("商务响应：交付周期 30 天，质保期 12 个月。")
    doc.add_paragraph("技术响应：电压等级 10kV，短路电流 30kA。")
    doc.save(buf)
    return buf.getvalue()


def make_xlsx() -> bytes:
    from openpyxl import Workbook

    buf = io.BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.append(["材料", "数量"])
    ws.append(["电缆", "100"])
    wb.save(buf)
    return buf.getvalue()


def make_pptx() -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"/>',
        )
        zf.writestr(
            "ppt/slides/slide1.xml",
            '<?xml version="1.0"?><p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" '
            'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><p:sp><p:txBody><a:p><a:r>'
            "<a:t>投标注意事项：履约保证金 5%，须在响应文件中单独承诺。</a:t></a:r></a:p></p:txBody></p:sp></p:sld>",
        )
    return buf.getvalue()


def make_ofd() -> bytes:
    ns = "http://www.ofdspec.org/2016"
    ofd_xml = (
        f'<?xml version="1.0" encoding="UTF-8"?><ofd:OFD xmlns:ofd="{ns}"><ofd:DocBody>'
        f"<ofd:DocRoot><ofd:BaseLoc>Doc_0/Document.xml</ofd:BaseLoc></ofd:DocRoot>"
        f"</ofd:DocBody></ofd:OFD>"
    )
    document_xml = (
        f'<?xml version="1.0" encoding="UTF-8"?><ofd:Document xmlns:ofd="{ns}"><ofd:Pages>'
        f'<ofd:Page ID="1"><ofd:BaseLoc>Doc_0/Pages/Page_1/Content.xml</ofd:BaseLoc></ofd:Page>'
        f"</ofd:Pages></ofd:Document>"
    )
    content_xml = (
        f'<?xml version="1.0" encoding="UTF-8"?><ofd:Page xmlns:ofd="{ns}"><ofd:Content>'
        f'<ofd:Layer ID="4" Type="Foreground"><ofd:TextObject ID="5" Font="3" Size="12" '
        f'Boundary="0 0 100 20"><ofd:TextCode X="0" Y="12">投标人须提供近三年同类业绩证明。</ofd:TextCode>'
        f"</ofd:TextObject></ofd:Layer></ofd:Content></ofd:Page>"
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("OFD.xml", ofd_xml)
        zf.writestr("Doc_0/Document.xml", document_xml)
        zf.writestr("Doc_0/Pages/Page_1/Content.xml", content_xml)
    return buf.getvalue()


def make_pdf() -> bytes | None:
    try:
        import fitz  # pymupdf
    except ImportError:
        return None
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "技术规范：主变容量 100MVA，电压等级 10kV。")
    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    return buf.getvalue()


def make_corrupt_docx() -> bytes:
    """合法 zip 容器（过 upload 校验）+ 垃圾内容（python-docx 解析必失败）。"""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("junk.bin", b"not an ooxml package at all")
    return buf.getvalue()


def upload(base: str, token: str, pid: int, name: str, data: bytes) -> dict:
    r = api(
        base,
        "POST",
        "/api/v1/files/upload",
        token=token,
        files={"files": (name, data, "application/octet-stream")},
        data={"target": "project", "project_id": str(pid)},
        timeout=90,
    )
    return r["files"][0]


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--base", default=BASE)
    parser.add_argument("--tag", default="seeded")
    args = parser.parse_args()
    base = args.base.rstrip("/")
    tag = args.tag

    for _ in range(30):
        try:
            if httpx.get(f"{base}/healthz", timeout=2).status_code == 200:
                break
        except Exception:  # noqa: BLE001
            time.sleep(1)
    else:
        print("API 未就绪")
        return 1

    stamp = int(time.time())
    email = f"seeded-{stamp}@test.com"
    pwd = "Abc12345"

    # ============ Phase 1：种子历史账号（API 层） ============
    reg = api(base, "POST", "/api/v1/auth/register", json={
        "email": email, "password": pwd, "enterprise_name": f"种子态E2E-{stamp % 100000}",
    })
    token = reg["access_token"]
    record("种子账号注册", True, email)

    pids = []
    for i in (1, 2, 3):
        p = api(base, "POST", "/api/v1/projects", token=token, json={"name": f"历史项目{i}"})
        pids.append(p["project_id"])
        ev = api(base, "POST", f"/api/v1/projects/{p['project_id']}/evaluate", token=token, json={}, timeout=240)
        record(f"项目{i} 评审产生评分", bool(ev.get("score_id")), f"score_id={ev.get('score_id')}")

    # 精确回归：多项目各有评分 → 列表 200 且各取最新评分（旧代码此处 PG 500）
    try:
        lst = api(base, "GET", "/api/v1/projects?size=50", token=token)
        by_id = {p["project_id"]: p for p in lst.get("items", [])}
        scores = [by_id[pid]["summary"]["latest_total_score"] for pid in pids if pid in by_id]
        ok = len(scores) == 3 and all(s is not None for s in scores)
        record("PG 多项目评分列表回归(API)", ok, f"latest_total_score={scores}")
    except httpx.HTTPStatusError as e:
        record("PG 多项目评分列表回归(API)", False, f"HTTP {e.response.status_code}")

    # ============ Phase 2：浏览器登录历史账号 → 列表 UI ============
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page(viewport={"width": 1440, "height": 900})
        page.goto(f"{base}/demo/")
        page.wait_for_load_state("networkidle")
        try:
            page.fill("#a-email", email)
            page.fill("#a-pwd", pwd)
            page.click("#a-login")
            ok_login = wait_log(page, "login 成功", timeout_ms=30000)
            record("历史账号浏览器登录", ok_login)
            rows_ok = page.wait_for_function(
                "() => { const el = document.getElementById('p-rows'); return el && el.querySelectorAll('tr').length >= 3; }",
                timeout=30000,
            )
            log_clean = "项目列表失败" not in log_text(page)
            record("历史账号列表 UI 回归(3 行无失败)", rows_ok and log_clean,
                   f"rows={page.evaluate('document.getElementById(\"p-rows\").querySelectorAll(\"tr\").length')} 失败日志={'有' if not log_clean else '无'}")
        except Exception as e:  # noqa: BLE001
            record("历史账号浏览器登录/列表", False, str(e)[:120])

        # ============ Phase 3：输入矩阵（API 上传 + 一条 UI 损坏文件） ============
        pid4 = api(base, "POST", "/api/v1/projects", token=token, json={"name": "输入矩阵项目"})["project_id"]
        valid_files = [
            ("招标矩阵.txt", make_txt()),
            ("商务技术矩阵.docx", make_docx()),
            ("清单矩阵.xlsx", make_xlsx()),
            ("注意事项矩阵.pptx", make_pptx()),
            ("业绩要求矩阵.ofd", make_ofd()),
        ]
        pdf = make_pdf()
        if pdf:
            valid_files.append(("技术规范矩阵.pdf", pdf))
        matrix_ok = True
        matrix_note = []
        for name, data in valid_files:
            try:
                item = upload(base, token, pid4, name, data)
                matrix_ok = matrix_ok and item.get("status") == 3
                matrix_note.append(f"{name.split('.')[-1]}:{item.get('status')}")
            except Exception as e:  # noqa: BLE001
                matrix_ok = False
                matrix_note.append(f"{name}:ERR {str(e)[:60]}")
        record(f"多格式输入矩阵({len(valid_files)} 种全部解析成功)", matrix_ok, " ".join(matrix_note))

        # 损坏 docx：上传即整体拒绝（红字原因），不入库不进资料列表（Issue #8 复盘语义）
        pid5 = api(base, "POST", "/api/v1/projects", token=token, json={"name": "损坏文件项目"})["project_id"]
        corrupt = upload(base, token, pid5, "坏文件.docx", make_corrupt_docx())
        reason = str(corrupt.get("error") or "")
        files_after = api(base, "GET", f"/api/v1/files?target=project&project_id={pid5}&size=100", token=token)
        not_in_list = all(x["name"] != "坏文件.docx" for x in files_after.get("items", []))
        record("损坏 docx 上传即拒绝且不入资料列表(API)", bool(reason) and "文件解析失败" in reason and not_in_list,
               f"reason={reason[:60]} 入列表={'是' if not not_in_list else '否'}")

        # 同一损坏文件走 UI 上传 → 上传即红字拒绝（模拟产品路径）
        try:
            tab(page, "项目")
            page.wait_for_function(
                "() => document.querySelectorAll('#p-rows .row-use').length >= 5", timeout=30000,
            )
            # 选用"损坏文件项目"（按名称定位）
            page.click(f"#p-rows tr:has-text('损坏文件项目') .row-use")
            page.wait_for_timeout(500)
            tab(page, "资料")
            page.set_input_files("#m-file", {
                "name": "坏文件UI.docx", "mimeType": "application/octet-stream",
                "buffer": make_corrupt_docx(),
            })
            page.click("button:has-text('上传')")
            ok_notice = wait_log(page, "上传被拒绝", timeout_ms=60000)
            notice_text = log_text(page)
            ok_reason = "文件解析失败" in notice_text and "未知错误" not in notice_text
            rows_empty = page.evaluate(
                "() => { const el = document.getElementById('m-files'); return el && el.querySelectorAll('tr').length === 0; }"
            )
            record("UI 上传损坏 docx 红字拒绝且列表为空", ok_notice and ok_reason and rows_empty,
                   f"原因可见={'文件解析失败' in notice_text} 未知错误={'未知错误' in notice_text} 列表空={rows_empty}")
        except Exception as e:  # noqa: BLE001
            record("UI 上传损坏 docx 红字拒绝且列表为空", False, str(e)[:120])

        # ============ Phase 4：错误路径任务（坏 file_id 触发失败，禁止"未知错误"） ============
        # 上传即拒绝语义下，任务级解析失败只剩存量坏数据等边界；用坏 file_id 精确触发失败路径，
        # 校验 API error 与 SSE failed 事件都携带具体原因。
        try:
            t = api(base, "POST", f"/api/v1/projects/{pid5}/tasks", token=token, json={
                "task_type": "tender_parse", "payload": {"file_ids": [999999]}, "idempotency_key": f"badfid-{stamp}",
            })
            task_id = t["task_id"]
            final = poll_task(base, token, task_id, 240)
            api_ok = final.get("status") == 6 and "文件不存在" in str(final.get("error") or "")
            record("失败任务 API error.message 含原因", api_ok, str(final.get("error"))[:80])
            try:
                r = httpx.get(f"{base}/api/v1/tasks/{task_id}/stream", headers={"Authorization": f"Bearer {token}"}, timeout=30)
                body = r.text
                sse_ok = "event: failed" in body and '"error"' in body and "文件不存在" in body
            except Exception as e:  # noqa: BLE001
                body, sse_ok = "", False
            record("SSE failed 事件携带 error", sse_ok, f"stream_len={len(body)}")
        except Exception as e:  # noqa: BLE001
            record("错误路径任务(坏 file_id)", False, str(e)[:120])

        browser.close()

    # ============ Phase 5：有效矩阵项目整体解析成功 ============
    fids = [upload(base, token, pid4, n, d)["file_id"] for n, d in valid_files]
    try:
        t = api(base, "POST", f"/api/v1/projects/{pid4}/tasks", token=token, json={
            "task_type": "tender_parse", "payload": {"file_ids": fids}, "idempotency_key": f"matrix-{stamp}",
        })
        final = poll_task(base, token, t["task_id"], 300)
        ok = final.get("status") == 3 and final.get("result") is not None
        record("矩阵项目整体 tender_parse 完成", ok, f"task={t['task_id']} status={final.get('status')}")
    except Exception as e:  # noqa: BLE001
        record("矩阵项目整体 tender_parse 完成", False, str(e)[:120])

    failed = [r for r in results if not r[1]]
    print(f"\n==== 种子态回归汇总：{len(results) - len(failed)}/{len(results)} PASS (base={base}, tag={tag}) ====", flush=True)
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
