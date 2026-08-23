"""DOCX 导出回归（Issue #8 + 产品要求：唯一生成路径=底稿式；全部改动以修订模式+批注记录）。"""

from __future__ import annotations

import io
import zipfile
from urllib.parse import unquote

from docx import Document

from app.services.export_service import docx_from_template


def _make_source() -> bytes:
    doc = Document()
    doc.add_paragraph("第一章 投标人须知")
    doc.add_paragraph("本采购由（采购人）实施，项目为（项目名称）。")
    doc.add_paragraph("供应商名称：（响应供应商名称）")
    doc.add_paragraph("联系电话：____")
    doc.add_paragraph("致：采购人")
    # 授权委托书形态：标签紧邻空位（下划线才是字段本身）
    doc.add_paragraph(
        "特授权____代表我方全权办理____（项目名称）（采购编号）（分标编号）（包号）"
        "项目的响应、谈判、签约、执行等具体工作。"
    )
    tb = doc.add_table(rows=1, cols=2)
    tb.rows[0].cells[0].text = "条款"
    tb.rows[0].cells[1].text = "说明"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _doc_xml(data: bytes):
    """从产物 zip 读 document.xml 并用纯 lxml 解析（Word 实际读取的序列化形态）。

    注意：不能用 python-docx 重开后的元素做 itertext——其 CT_P/CT_R 的 .text 是
    计算属性，逐层重复取值会让同一文本出现多份（读取侧假象，序列化产物无此问题）。"""
    import zipfile

    from lxml import etree

    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        return etree.fromstring(zf.read("word/document.xml"))


def _all_text(data: bytes) -> str:
    return "".join(_doc_xml(data).itertext())


def _count_el(data: bytes, local_tag: str) -> int:
    return sum(1 for e in _doc_xml(data).iter() if e.tag.split("}")[-1] == local_tag)


def _comment_texts(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        names = zf.namelist()
        assert "word/comments.xml" in names, f"缺少批注部件，包内: {names[:10]}"
        return zf.read("word/comments.xml").decode("utf-8")


def test_draft_export_keeps_whole_source_and_fills_tracked(tmp_path):
    src = tmp_path / "source.docx"
    src.write_bytes(_make_source())
    model = {
        "buyer": "测试招标人",
        "project_name": "测试采购项目",
        "supplier_name": "测试供应商",
        "supplement_nodes": [
            {"type": "heading", "text": "一、总体响应说明"},
            {"type": "paragraph", "text": "我方完全响应采购文件全部要求。"},
        ],
    }
    data = docx_from_template(str(src), model)
    joined = _all_text(data)
    doc = Document(io.BytesIO(data))
    # 整本保留：原段落都在，表格保留
    assert "第一章 投标人须知" in joined
    assert len(doc.tables) == 1
    # 已知字段回填（修订插入可见）
    assert "本采购由测试招标人实施，项目为测试采购项目。" in joined
    assert "供应商名称：测试供应商" in joined
    assert "联系电话：【待补充：联系电话】" in joined
    # 应答函裸标签形态：致：采购人 → 致：真实招标人
    assert "致：测试招标人" in joined
    # 标签紧邻空位整体回填：____（项目名称）→ 项目名，而不是【待补充】+项目名
    assert "全权办理测试采购项目（采购编号）" in joined
    assert "【待补充】测试采购项目" not in joined
    # 补充节追加
    assert "补充响应内容" in joined
    assert "我方完全响应采购文件全部要求。" in joined
    # 修订模式：删除（原文删除线）与插入标记、批注引用
    assert _count_el(data, "ins") > 0
    assert _count_el(data, "del") > 0
    assert _count_el(data, "commentReference") > 0
    # Word 显示修订的前提：settings.xml 含 trackChanges
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        settings = zf.read("word/settings.xml").decode("utf-8")
    assert "<w:trackChanges" in settings
    # 批注部件存在且说明来源
    comments = _comment_texts(data)
    assert "来源" in comments and "BidVolt" in comments
    # 批注范围顺序：commentRangeStart → del → ins → commentRangeEnd → commentReference
    root = _doc_xml(data)
    for p in root.iter():
        if p.tag.split("}")[-1] != "p":
            continue
        tags = [ch.tag.split("}")[-1] for ch in p if isinstance(ch.tag, str)]
        if "commentRangeStart" not in tags:
            continue
        i_start = tags.index("commentRangeStart")
        i_del = tags.index("del")
        i_ins = tags.index("ins")
        i_end = tags.index("commentRangeEnd")
        i_ref = next(i for i, t in enumerate(tags) if t == "r")
        assert i_start < i_del < i_ins < i_end < i_ref, tags


def test_draft_export_fills_unknown_with_placeholder(tmp_path):
    src = tmp_path / "source.docx"
    src.write_bytes(_make_source())
    model = {"buyer": "", "project_name": "", "supplier_name": "", "supplement_nodes": []}
    data = docx_from_template(str(src), model)
    joined = _all_text(data)
    assert "【待补充：招标人名称】" in joined
    assert "【待补充：项目名称】" in joined
    assert "【待补充：供应商名称】" in joined


def test_draft_export_supplement_headings_bold_without_style_dependency(tmp_path):
    """追加标题不依赖源文档是否含内置 Heading 样式（国网模板曾因此 KeyError 回退）。"""
    src = tmp_path / "source.docx"
    src.write_bytes(_make_source())
    model = {
        "buyer": "",
        "project_name": "",
        "supplier_name": "",
        "supplement_nodes": [
            {"type": "heading", "text": "补充节标题"},
            {"type": "paragraph", "text": "（采购人）应对本项目负责。"},
        ],
    }
    data = docx_from_template(str(src), model)
    joined = _all_text(data)
    assert "【待补充：招标人名称】应对本项目负责。" in joined
    # 标题加粗（w:b 存在）
    assert _count_el(data, "b") > 0
    # 补充节内容标记为插入
    assert _count_el(data, "ins") > 0


def test_split_supplement_tolerates_string_nodes():
    """撰写内容节点兼容裸字符串（新方案 agent 落库形状）与混合形状。"""
    from app.services.export_service import _split_supplement

    titles = ["（一）响应函及报价汇总表", "（二）报价明细表"]
    nodes = [
        {"type": "heading", "text": "（一）响应函及报价汇总表"},
        "我方承诺……（字符串节点）",
        {"type": "heading", "text": "（二）报价明细表"},
        {"type": "paragraph", "text": "报价明细"},
    ]
    matched, rest = _split_supplement(titles, nodes)
    assert len(matched[0]) == 2  # heading + 字符串节点（归一到 paragraph）分配到第 0 条
    assert len(matched[1]) == 2  # heading + paragraph 分配到第 1 条
    assert rest == []


def test_labeled_space_blank_fill():
    """带标签空位（空格/零宽形态）：有资料回填、无资料原位【待补充】，全程修订模式。"""
    from docx import Document

    src = Document()
    src.add_paragraph("1. 我方已仔细研究了采购编号：    ，项目名称：   ，分标名称：   ，包名称：   采购文件的全部内容。")
    src.add_paragraph("响应供应商：")
    src.add_paragraph("响应供应商：（盖章）")
    src.add_paragraph("电话：")
    src.add_paragraph("法定地址：。")
    src.add_paragraph("附表1 项目名称：，采购编号：，包号：，单位：万元人民币")
    buf = io.BytesIO()
    src.save(buf)
    p = __import__("pathlib").Path(__import__("tempfile").gettempdir()) / "label_blank_src.docx"
    p.write_bytes(buf.getvalue())

    from app.services.export_service import _FillSession

    doc = Document(str(p))
    sess = _FillSession(doc, "采购人甲", "项目乙", "供应商丙", "412623-1")
    sess.apply_to_doc()
    data = sess.finish()

    joined = _all_text(data)
    assert "采购编号：412623-1，项目名称：项目乙，分标名称：【待补充：分标名称】，包名称：【待补充：包名称】采购文件的全部内容" in joined
    assert "响应供应商：供应商丙" in joined
    assert "响应供应商：供应商丙（盖章）" in joined
    assert "电话：【待补充：电话】" in joined
    assert "法定地址：【待补充：法定地址】" in joined
    assert "项目名称：项目乙，采购编号：412623-1，包号：【待补充：包号】，单位：万元人民币" in joined
    # 修订模式 + 批注（改什么都留痕）
    assert _count_el(data, "ins") > 0 and _count_el(data, "del") > 0
    assert "采购编号=「412623-1」" in _comment_texts(data)
    # 原文保留：修订删除线里仍有模板原句（空格位原文）——删除线数量>0 即可（空格删除也留痕）
    assert _count_el(data, "delText") > 0


def test_assembly_primitives_roundtrip(tmp_path):
    """成文工具链机制原语：切片→填空→校验→封存 一轮闭环（旧路径不动）。"""
    from docx import Document
    from docx.oxml.ns import qn

    from app.services.export_service import (
        _build_item_document,
        _FillSession,
        check_doc_fidelity,
        replace_text_tracked,
    )
    src = Document()
    src.add_paragraph("（一）响应函及报价汇总表")
    src.add_paragraph("1. 我方已仔细研究了采购编号：    ，项目名称：   ，包名称：   采购文件。")
    p = tmp_path / "src.docx"
    src.save(str(p))

    src2 = Document(str(p))
    elems = [("p", el) for el in src2.element.body.iterchildren() if el.tag == qn("w:p")]
    doc, located = _build_item_document(str(p), None, elems)
    assert located is True

    # 填空（标准字段 + 显式定向替换）
    sess = _FillSession(doc, "采购人甲", "项目乙", "供应商丙", "412623-1")
    sess.apply_to_doc()
    n = replace_text_tracked(sess.editor, "包名称：【待补充：包名称】", "包名称：包A", "主会话指定应答分包")
    assert n == 1

    # 校验：原文⊂底稿必须通过（删除线保留原文，插入不算改写）
    import zipfile as _z

    from lxml import etree as _e

    with _z.ZipFile(str(p)) as zf:
        src_text = "".join(_e.fromstring(zf.read("word/document.xml")).itertext())
    r = check_doc_fidelity(doc, src_text)
    assert r["ok"] is True, r["issues"]

    data = sess.finish()
    joined = _all_text(data)
    assert "采购编号：412623-1" in joined
    assert "项目名称：项目乙" in joined
    assert "包名称：包A" in joined
    assert _count_el(data, "ins") > 0 and _count_el(data, "del") > 0


def test_locate_item_elements_matches_requested_item(tmp_path):
    """同部分多条目的切片必须按标题匹配请求条目（回归：曾对每份条目都返回首条目区间，
    致 价格/商务/技术 各文件内容雷同）。清单标题与底稿标题的编号/括号注差异也要容错。"""
    from types import SimpleNamespace

    from docx import Document
    from docx.oxml.ns import qn

    from app.services.export_service import locate_item_elements, locate_item_with_heading

    def add_with_lvl(doc, text, lvl):
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        pPr.append(pPr.makeelement(qn("w:outlineLvl"), {qn("w:val"): str(lvl)}))
        return p

    src = Document()
    add_with_lvl(src, "第五章 响应文件格式", 0)
    add_with_lvl(src, "商务文件", 1)
    add_with_lvl(src, "（一）法定代表人（单位负责人）授权委托书", 2)
    src.add_paragraph("授权委托书正文：兹授权……")
    add_with_lvl(src, "（二）商务偏差表", 2)
    src.add_paragraph("偏差表正文：无偏差。")
    add_with_lvl(src, "（三）响应保证保险（如有）", 2)
    src.add_paragraph("保险正文：保单附后。")
    src.add_paragraph("第六章 商务评分标准")
    p = tmp_path / "src.docx"
    src.save(str(p))

    def row(req_id, title):
        return SimpleNamespace(id=req_id, content=title, structured={"role": "business"})

    def text_of(elems):
        return "".join("".join(el.itertext()) for _k, el in elems)

    e1 = locate_item_elements(str(p), row(1, "（一）法定代表人授权委托书"))
    assert e1 is not None
    t1 = text_of(e1)
    assert "授权委托书正文" in t1 and "商务偏差表" not in t1

    e2 = locate_item_elements(str(p), row(2, "（二）商务偏差表"))
    assert e2 is not None
    t2 = text_of(e2)
    assert "偏差表正文" in t2 and "授权委托书正文" not in t2 and "保险正文" not in t2

    e3 = locate_item_elements(str(p), row(3, "（三）响应保证保险（如有）"))
    assert e3 is not None
    t3 = text_of(e3)
    assert "保险正文" in t3 and "偏差表正文" not in t3

    # 不存在的条目：不串到别的条目，返回 None 走清单重建+批注路径
    assert locate_item_elements(str(p), row(4, "（四）补充文件")) is None

    # 身份信号：matched_heading 必须等于实际绑定的底稿条目标题（供调用方比对防串区）
    h1, e1b = locate_item_with_heading(str(p), row(1, "（一）法定代表人授权委托书"))
    assert e1b is not None and "法定代表人（单位负责人）授权委托书" in h1
    h2, e2b = locate_item_with_heading(str(p), row(2, "（二）商务偏差表"))
    assert e2b is not None and h2.startswith("（二）商务偏差表")
    h3, e3b = locate_item_with_heading(str(p), row(3, "（三）响应保证保险（如有）"))
    assert e3b is not None and "响应保证保险" in h3
    h4, e4b = locate_item_with_heading(str(p), row(4, "（四）补充文件"))
    assert e4b is None and h4 is None


def test_fidelity_ignores_drawing_internals(tmp_path):
    """锚定图（印模等浮动图片）的坐标数字（wp:posOffset/extent）不是正文：
    底稿侧与条目侧必须用同一收集器（canonical_text），否则合法切片被误报不忠实。
    （回归：补充文件切片因底稿 itertext 混入 wp:posOffset 数字而 verify 失败。）"""
    import copy

    from lxml import etree

    from app.services.export_service import canonical_text, check_doc_fidelity

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    W_NS = f"{{{W}}}"
    WP_NS = f"{{{WP}}}"

    def w(tag):
        return W_NS + tag

    # 构造段落：正文 + 锚定图（wp:posOffset 数字噪声）
    p = etree.Element(w("p"), nsmap={"w": W, "wp": WP})
    r = etree.SubElement(p, w("r"))
    t = etree.SubElement(r, w("t"))
    t.text = "响应供应商投标专用章"
    anchor = etree.SubElement(r, WP_NS + "anchor")
    po = etree.SubElement(anchor, WP_NS + "posOffset")
    po.text = "285750"
    ext = etree.SubElement(anchor, WP_NS + "extent")
    ext.set("cx", "2066925")

    doc_el = etree.Element(w("document"), nsmap={"w": W})
    body = etree.SubElement(doc_el, w("body"))
    body.append(p)

    src_text = canonical_text(doc_el)
    assert src_text == "响应供应商投标专用章", src_text
    assert "285750" not in src_text and "2066925" not in src_text

    # 切片=同段落深拷贝：忠实性必须通过（两侧同算法）
    doc = Document()
    doc.element.body.append(copy.deepcopy(p))
    r = check_doc_fidelity(doc, src_text)
    assert r["ok"] is True, r["issues"]
    assert r["original_chars"] == len("响应供应商投标专用章")

    # 修订插入（w:ins）同样不计原文（两侧一致）
    ins = etree.SubElement(p, w("ins"))
    ir = etree.SubElement(ins, w("r"))
    it = etree.SubElement(ir, w("t"))
    it.text = "北京北辰电力科技有限公司"
    assert "北京北辰" not in canonical_text(doc_el)
    doc2 = Document()
    doc2.element.body.append(copy.deepcopy(p))
    r2 = check_doc_fidelity(doc2, src_text)
    assert r2["ok"] is True
    assert r2["inserted_chars"] == len("北京北辰电力科技有限公司")


def test_table_fill_in_place():
    """表格/表单类条目：内容填进模板自身表格单元格（修订插入+批注），而不是空表挂文末。"""
    from docx import Document

    from app.services.export_service import _FillSession

    src = Document()
    tb = src.add_table(rows=2, cols=2)
    tb.rows[0].cells[0].text = "项目名称"
    tb.rows[0].cells[1].text = "项目规模"
    src.add_paragraph("（一）响应函及报价汇总表")

    sess = _FillSession(src, "采购人甲", "项目乙", "供应商丙", "412623-1")
    assert sess.fill_table_cell(0, 1, 0, "虚拟电厂数据融合系统", "业绩表第1行填写") is True
    assert sess.fill_table_cell(0, 1, 1, "覆盖浙江全境", None) is True
    assert sess.fill_table_cell(5, 0, 0, "x") is False  # 越界如实返回

    data = sess.finish()
    joined = _all_text(data)
    assert "虚拟电厂数据融合系统" in joined
    assert "覆盖浙江全境" in joined
    assert _count_el(data, "ins") > 0
    assert "业绩表第1行填写" in _comment_texts(data)


def test_append_supplement_heading_param():
    """追加节标题由调用方（主会话）按投标文体自定，默认不再带系统痕迹字样。"""
    from docx import Document

    from app.services.export_service import _FillSession

    src = Document()
    src.add_paragraph("模板正文")
    sess = _FillSession(src, "", "", "", "")
    sess.append_supplement(
        [{"type": "paragraph", "text": "方案正文内容"}],
        heading_text="一、对项目的理解",
        comment=None,
        page_break=False,
    )
    data = sess.finish()
    joined = _all_text(data)
    assert "一、对项目的理解" in joined
    assert "方案正文内容" in joined
    assert "主会话撰写" not in joined


def test_labeled_blanks_fill_enterprise_facts():
    """带标签空位：agent 经通用 label_values 给任意标签填值（含清单外标签），不落【待补充】。"""
    from docx import Document

    from app.services.export_service import _FillSession

    src = Document()
    src.add_paragraph("单位地址：，法定地址：。")
    src.add_paragraph("法定代表人（单位负责人）或授权代表：，电话：，邮政编码：，传真：。")
    src.add_paragraph("分标名称：，标段名称：。")  # 标段名称=清单外标签，通用机制照样填
    sess = _FillSession(
        src, "采购人甲", "项目乙", "供应商丙", "412623-1",
        label_values={
            "单位地址": "北京市海淀区示例路1号",
            "法定地址": "北京市海淀区示例路1号",
            "法定代表人（单位负责人）或授权代表": "张建国",
            "电话": "010-88886666",
            "邮政编码": "100000",
            "传真": "010-88886667",
            "分标名称": "标段一",
            "标段名称": "标段一（自定义标签）",
        },
    )
    sess.apply_to_doc()
    data = sess.finish()
    joined = _all_text(data)
    assert "单位地址：北京市海淀区示例路1号" in joined
    assert "法定代表人（单位负责人）或授权代表：张建国" in joined
    assert "电话：010-88886666" in joined
    assert "邮政编码：100000" in joined
    assert "分标名称：标段一" in joined
    assert "标段名称：标段一（自定义标签）" in joined
    # 无资料才【待补充】
    sess2 = _FillSession(Document(), "", "", "", "")
    assert "【待补充：电话】" in sess2._labeled_value("电话")


def test_remaining_blanks_signal():
    """填完即反馈：remaining_blanks 逐处列出本文档还有哪些空位没填（标签+上下文）。"""
    from docx import Document

    from app.services.export_service import _FillSession

    src = Document()
    src.add_paragraph("电话：，分标名称：，单位地址：。")
    sess = _FillSession(src, "", "", "", "", label_values={"电话": "010-88886666"})
    sess.apply_to_doc()
    items = sess.remaining_blanks()
    labels = [i["label"] for i in items]
    assert "分标名称" in labels and "单位地址" in labels
    assert "电话" not in labels  # 已填的不出现在剩余清单
    ctx = {i["label"]: i.get("context", "") for i in items}
    assert "分标" in ctx["分标名称"] or ctx["分标名称"] == ""


def test_underscore_blanks_auto_labeled():
    """裸下划线空位：前文有"标签："自动带标签，无标签上下文才裸【待补充】。"""
    from docx import Document

    from app.services.export_service import _FillSession

    src = Document()
    src.add_paragraph("不含税单价：________，税率：13%。")
    src.add_paragraph("特授权________代表我方全权办理。")  # 无标签上下文 → 裸（留给 agent fills）
    sess = _FillSession(src, "", "", "", "")
    sess.apply_to_doc()
    data = sess.finish()
    joined = _all_text(data)
    assert "不含税单价：【待补充：不含税单价】" in joined
    assert "特授权【待补充】代表我方全权办理" in joined


def test_draft_label_format():
    from app.services.export_service import _draft_label

    assert _draft_label("采购文件.docx", "商务标") == "商务标(底稿：采购文件.docx)"
    assert _draft_label(None, "商务标") == "商务标"


def test_xlsx_bytes_shape_tolerant():
    """xlsx_bytes 兼容旧 {sheets} 与新方案 agent 落库的 {docModel:{sections}} 形状，
    任何形状下都必须产出至少一个可见工作表（openpyxl 硬性要求）。"""
    from openpyxl import load_workbook

    from app.services.export_service import xlsx_bytes

    def sheets_of(data):
        wb = load_workbook(io.BytesIO(data))
        assert len(wb.worksheets) >= 1, "必须至少有一个可见工作表"
        return wb

    # 旧形状：显式 sheets
    wb = sheets_of(xlsx_bytes({"sheets": [{"name": "报价单", "rows": [["a", "1"]]}]}))
    assert wb.active["A1"].value == "a"

    # 新方案形状：docModel.sections（agent 主会话落库形态）
    wb = sheets_of(
        xlsx_bytes(
            {
                "docModel": {
                    "title": "价格标",
                    "sections": [
                        {"title": "1. 响应函", "content": "我方承诺。/n 第二条。", "source": "block 1~9"},
                    ],
                }
            }
        )
    )
    ws = wb.active
    assert ws["A1"].value == "条目"
    assert "我方承诺。\n 第二条。" in str(ws["B2"].value)

    # 空模型 / 完全无表数据：单可见工作表 + 说明行，不得 500
    wb = sheets_of(xlsx_bytes({}))
    assert wb.active["A1"].value is not None

    # 被 {version_no, model} 包装的形状
    wb = sheets_of(xlsx_bytes({"version_no": 2, "model": {"rows": [["x", "y"]]}}))
    assert wb.active["A1"].value == "x"

    # 生成侧序列化退化：rows 写成对象（{"item": [...]}）必须报可操作的错，不得静默产出残缺表
    import pytest

    with pytest.raises(ValueError, match="数组的数组"):
        xlsx_bytes({"sheets": [{"name": "报价单", "rows": {"item": ["合计"]}}]})


def test_draft_download_filename_marks_source(client):
    """下载文件名标注底稿来源（产品要求：一眼可见底稿是谁）。"""
    r = client.post(
        "/api/v1/auth/register",
        json={"email": "fn@test.com", "password": "Abc12345", "enterprise_name": "文件名企业"},
    )
    h = {"Authorization": f"Bearer {r.json()['access_token']}"}
    pid = client.post("/api/v1/projects", json={"name": "P"}, headers=h).json()["project_id"]

    src = Document()
    src.add_paragraph("采购文件原文。")
    buf = io.BytesIO()
    src.save(buf)
    up = client.post(
        "/api/v1/files/upload",
        data={"target": "project", "project_id": str(pid)},
        files=[("files", ("采购文件.docx", buf.getvalue(), "application/octet-stream"))],
        headers=h,
    )
    assert up.status_code == 200

    did = client.post(
        "/api/v1/deliverables",
        json={"project_id": pid, "deliverable_type": 1, "title": "商务标"},
        headers=h,
    ).json()["deliverable_id"]
    client.post(
        f"/api/v1/deliverables/{did}/versions",
        json={"content": {"nodes": [{"id": "n1", "type": "paragraph", "text": "商务响应"}]}},
        headers=h,
    )
    r = client.get(f"/api/v1/deliverables/{did}/versions/1/download", headers=h)
    assert r.status_code == 200
    disposition = unquote(r.headers["content-disposition"])
    assert "底稿：采购文件.docx" in disposition, disposition
