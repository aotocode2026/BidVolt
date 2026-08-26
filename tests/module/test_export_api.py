from __future__ import annotations

import asyncio
import io
import zipfile

from sqlalchemy import select
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

from app.models.auth import AppUser

TEST_DB = "./.test_bidvolt.db"


def _setup(client):
    r = client.post(
        "/api/v1/auth/register",
        json={"email": "x@test.com", "password": "Abc12345", "enterprise_name": "测试企业"},
    )
    headers = {"Authorization": f"Bearer {r.json()['access_token']}"}
    pid = client.post("/api/v1/projects", json={"name": "P"}, headers=headers).json()["project_id"]
    return headers, pid


def _add_deliverable(client, headers, pid, dtype, model):
    did = client.post(
        "/api/v1/deliverables",
        json={"project_id": pid, "deliverable_type": dtype, "title": "成果"},
        headers=headers,
    ).json()["deliverable_id"]
    client.post(
        f"/api/v1/deliverables/{did}/versions",
        json={"content": model, "version_type": 2},
        headers=headers,
    )
    return did


def _real_nodes(prefix: str) -> dict:
    """终检质量门禁（Issue #12）要求正文 >=100 字：构造真实长度正文。"""
    body = (
        f"{prefix}正式投标正文。" + "本公司具备相应资质与业绩，人员设备资金保障到位，"
        "质量保证体系健全，售后服务响应及时，工期承诺满足招标要求，"
        "报价测算依据充分，愿承担相应法律责任。本公司具备相应资质与业绩，人员设备资金保障到位，"
        "质量保证体系健全，售后服务响应及时，工期承诺满足招标要求。"
    )
    return {"nodes": [{"id": "n1", "type": "heading", "text": prefix}, {"id": "n2", "type": "paragraph", "text": body}]}


def test_final_check_detects_missing_deliverable(client):
    h, pid = _setup(client)
    _add_deliverable(client, h, pid, 1, _real_nodes("商务标"))
    _add_deliverable(client, h, pid, 3, {"type": "sheet", "sheets": [{"name": "报价单", "rows": [["a", "1"]]}]})

    check = client.post(f"/api/v1/projects/{pid}/check", json={}, headers=h)
    assert check.status_code == 200
    assert check.json()["passed"] is False
    assert any("技术标" in i["message"] for i in check.json()["issues"])

    _add_deliverable(client, h, pid, 2, _real_nodes("技术标"))
    check2 = client.post(f"/api/v1/projects/{pid}/check", json={}, headers=h)
    assert check2.json()["passed"] is True

    got = client.get(f"/api/v1/projects/{pid}/check/{check2.json()['check_id']}", headers=h)
    assert got.json()["passed"] is True


def test_final_check_rejects_stub_and_markdown(client):
    """Issue #12：终检必须识别占位草稿与 Markdown 残留。"""
    h, pid = _setup(client)
    _add_deliverable(client, h, pid, 2, _real_nodes("技术标"))
    _add_deliverable(client, h, pid, 3, {"type": "sheet", "sheets": [{"name": "报价单", "rows": [["a", "1"]]}]})
    stub = {
        "nodes": [
            {"id": "n1", "type": "paragraph", "text": "一、企业基本情况（草稿由 BidVolt 确定性生成，待人工校核）"},
        ]
    }
    _add_deliverable(client, h, pid, 1, stub)
    check = client.post(f"/api/v1/projects/{pid}/check", json={}, headers=h)
    assert check.json()["passed"] is False
    messages = [i["message"] for i in check.json()["issues"]]
    assert any("占位草稿" in m for m in messages)


def test_final_check_structure_compliance(client):
    """结构合规（用户反馈）：招标文件要求的章节缺失时必须拦截。"""
    from app.services.export_service import run_final_check

    class D:
        def __init__(self, dtype, vno):
            self.deliverable_type = dtype
            self.current_version_no = vno
            self.id = 100 + dtype

    long_text = "本公司具备相应资质与业绩，人员设备资金保障到位，质量保证体系健全，售后服务响应及时。"
    contents = {
        101: {"model": {"nodes": [{"type": "heading", "text": "一、应答函"},
                                {"type": "paragraph", "text": long_text}]}},
        102: {"model": {"nodes": [{"type": "heading", "text": "一、技术参数响应表"},
                                {"type": "paragraph", "text": long_text}]}},
        103: {"model": {"sheets": [{"name": "报价单", "rows": [["项目", "建议价"], ["x", "119.2"]]}]}},
    }
    result = run_final_check(
        [D(1, 1), D(2, 1), D(3, 1)],
        contents=contents,
        structure=[
            {"role": "business", "title": "一、应答函"},
            {"role": "technical", "title": "一、技术参数响应表"},
            {"role": "technical", "title": "二、技术方案"},  # 缺失
        ],
    )
    assert result["passed"] is False
    assert any(i["type"] == "结构合规" and "二、技术方案" in i["message"] for i in result["issues"])


def test_final_check_requirement_coverage_and_text_quality(client):
    """终检 v2：逐条要求覆盖（对标的要求）+ 重复段落 + 待补充占位 + 字数统计。"""
    from app.services.export_service import run_final_check

    class D:
        def __init__(self, dtype, vno):
            self.deliverable_type = dtype
            self.current_version_no = vno
            self.id = 200 + dtype

    long_para = "本公司具备相应资质与业绩，人员设备资金保障到位，质量保证体系健全，售后服务响应及时。"
    contents = {
        201: {"model": {"nodes": [{"type": "heading", "text": "商务标"}, {"type": "paragraph", "text": long_para}]}},
        202: {"model": {"nodes": [
            {"type": "heading", "text": "技术标"},
            {"type": "paragraph", "text": "电压等级 10kV 满足。"},
            {"type": "paragraph", "text": long_para},
            {"type": "paragraph", "text": long_para},  # 重复段落
            {"type": "paragraph", "text": "试验数据【待补充】"},
        ]}},
        203: {"model": {"sheets": [{"name": "报价单", "rows": [["项目", "建议价"], ["x", "119.2"]]}]}},
    }
    result = run_final_check(
        [D(1, 1), D(2, 1), D(3, 1)],
        requirements=[
            {"id": 1, "req_type": "tech_requirement", "content": "抗短路能力 30kA"},
            {"id": 2, "req_type": "qualification", "content": "资质能力保障到位"},
        ],
        contents=contents,
    )
    assert result["passed"] is False
    messages = [(i["type"], i["message"]) for i in result["issues"]]
    assert any(t == "要求覆盖" and "抗短路能力" in m for t, m in messages)  # 技术标未响应
    assert any(t == "要求覆盖" and "资质能力保障到位" in m for t, m in messages)  # 商务标未响应（前10字不在正文）
    assert any(t == "文字质量" and "重复段落" in m for t, m in messages)
    assert any(t == "文字质量" and "待补充" in m for t, m in messages)
    assert result["words"]["技术标"] > 0
    assert result["pending"]["技术标"] == 1


def test_export_and_delivery_package(client):
    h, pid = _setup(client)
    # 底稿式导出（唯一路径）：项目里必须有采购文件 docx 作底稿；无底稿时导出明确报错
    from docx import Document as _Docx

    src = _Docx()
    src.add_paragraph("采购文件原文：投标人须知。")
    src.add_paragraph("（项目名称）由（采购人）实施。")
    buf = io.BytesIO()
    src.save(buf)
    up = client.post(
        "/api/v1/files/upload",
        data={"target": "project", "project_id": str(pid)},
        files=[("files", ("采购文件.docx", buf.getvalue(), "application/octet-stream"))],
        headers=h,
    )
    assert up.status_code == 200
    assert up.json()["files"][0]["status"] == 3

    _add_deliverable(client, h, pid, 1, {"nodes": [{"id": "n1", "text": "商务响应"}]})
    _add_deliverable(client, h, pid, 2, {"nodes": [{"id": "n2", "text": "技术方案"}]})
    _add_deliverable(
        client,
        h,
        pid,
        3,
        {"type": "sheet", "sheets": [{"name": "报价单", "rows": [["材料", "价格"], ["电缆", "120"]]}]},
    )

    exp = client.post(
        f"/api/v1/projects/{pid}/export",
        json={"formats": ["docx", "xlsx"], "with_manifest": True},
        headers=h,
    )
    assert exp.status_code == 200
    job_id = exp.json()["job_id"]
    files = exp.json()["files"]
    assert any(f["name"].endswith(".docx") for f in files)
    assert any(f["name"].endswith(".xlsx") for f in files)
    assert any(f["name"] == "manifest.json" for f in files)
    assert all(f["sha256"] for f in files)

    # 底稿式：下载 docx 是整本底稿 + 填空 + 补充节（不是节点重排版）
    docx_entry = next(f for f in files if f["name"].endswith(".docx"))
    from app.services.storage import StorageProvider

    p = StorageProvider().open(docx_entry["bucket"], docx_entry["object_key"])
    out = _Docx(str(p))
    texts = [par.text for par in out.paragraphs]
    assert any("采购文件原文：投标人须知。" in t for t in texts)

    status = client.get(f"/api/v1/projects/{pid}/export/{job_id}", headers=h)
    assert status.json()["status"] == 2

    pkg = client.get(f"/api/v1/projects/{pid}/delivery-package", headers=h)
    assert pkg.status_code == 200
    with zipfile.ZipFile(io.BytesIO(pkg.content)) as zf:
        names = zf.namelist()
        assert "manifest.json" in names
        assert any(n.endswith(".docx") for n in names)
        assert any(n.endswith(".xlsx") for n in names)


def test_export_without_template_source_reports_clear_error(client):
    """产品要求：没有底稿就明确报错，绝不回退节点式生成。"""
    h, pid = _setup(client)
    _add_deliverable(client, h, pid, 1, {"nodes": [{"id": "n1", "text": "商务"}]})
    r = client.post(f"/api/v1/projects/{pid}/export", json={"formats": ["docx"]}, headers=h)
    assert r.status_code == 409
    assert "底稿" in r.json()["detail"]


def test_response_package_by_item_list(client):
    """产品定版：按招标文件《响应文件格式》清单逐份成文，装进三个目录。"""
    h, pid = _setup(client)
    # 底稿：含"响应文件格式"章 + 商务/技术两部分条目
    from docx import Document as _Docx

    src = _Docx()
    src.add_paragraph("第五章 响应文件格式")
    src.add_paragraph("响应文件由三部分组成：价格文件、商务文件、技术文件。")
    src.add_paragraph("商务文件")
    src.add_paragraph("（一）响应函")
    src.add_paragraph("致：（采购人）")
    src.add_paragraph("（二）商务偏差表")
    tb = src.add_table(rows=2, cols=3)
    tb.rows[0].cells[0].text = "序号"
    tb.rows[0].cells[1].text = "条目"
    tb.rows[0].cells[2].text = "响应"
    tb.rows[1].cells[0].text = "1"
    tb.rows[1].cells[1].text = "供货期"
    tb.rows[1].cells[2].text = "满足"
    src.add_paragraph("技术文件")
    src.add_paragraph("（一）技术偏差表")
    src.add_paragraph("技术偏差内容")
    src.add_paragraph("响应文件编制注意事项")
    buf = io.BytesIO()
    src.save(buf)
    up = client.post(
        "/api/v1/files/upload",
        data={"target": "project", "project_id": str(pid)},
        files=[("files", ("采购文件.docx", buf.getvalue(), "application/octet-stream"))],
        headers=h,
    )
    assert up.status_code == 200

    # 模板清单（响应文件格式逐字落库）
    r = client.post(
        f"/api/v1/projects/{pid}/requirements/upsert",
        json={"requirements": [
            {"req_type": "doc_template", "content": "（一）响应函",
             "structured": {"role": "business", "order": 1, "kind": "paragraph"}, "coordinates": [{"file_id": 1}]},
            {"req_type": "doc_template", "content": "（二）商务偏差表",
             "structured": {"role": "business", "order": 2, "kind": "table",
                            "rows": [["序号", "条目", "响应"], ["1", "供货期", "满足"]]}, "coordinates": [{"file_id": 1}]},
            {"req_type": "doc_template", "content": "（一）技术偏差表",
             "structured": {"role": "technical", "order": 1, "kind": "paragraph"}, "coordinates": [{"file_id": 1}]},
        ]},
        headers=h,
    )
    assert r.status_code == 201

    _add_deliverable(client, h, pid, 1, {
        "buyer": "测试招标人", "project_name": "测试采购项目", "supplier_name": "测试供应商",
        "supplement_nodes": [
            {"type": "heading", "text": "应答函"},
            {"type": "paragraph", "text": "我方已仔细研究采购文件全部内容。"},
            {"type": "heading", "text": "三、商务偏离表与承诺"},
            {"type": "paragraph", "text": "无偏离声明。"},
        ],
    })
    _add_deliverable(client, h, pid, 3, {"type": "sheet", "sheets": [{"name": "报价单", "rows": [["a", "1"]]}]})

    pkg = client.get(f"/api/v1/projects/{pid}/response-package", headers=h)
    assert pkg.status_code == 200
    with zipfile.ZipFile(io.BytesIO(pkg.content)) as zf:
        names = set(zf.namelist())
        assert "商务文件/（一）响应函.docx" in names, names
        assert "商务文件/（二）商务偏差表.docx" in names
        assert "技术文件/（一）技术偏差表.docx" in names
        assert "价格文件/报价单.xlsx" in names
        assert "manifest.json" in names
    # 响应函 = 底稿切片 + 填空 + 对应撰写内容
    from lxml import etree

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def _wt(root):
        return "".join(t.text or "" for t in root.iter(W + "t"))

    with zipfile.ZipFile(io.BytesIO(pkg.content)) as zf:
        inner = zipfile.ZipFile(io.BytesIO(zf.read("商务文件/（一）响应函.docx")))
        root = etree.fromstring(inner.read("word/document.xml"))
        text = _wt(root)  # 最终文本（含插入、不含删除层——最小差异修订）
        assert "致：测试招标人" in text  # 填空（修订插入）
        assert "我方已仔细研究采购文件全部内容。" in text  # 对应撰写内容已分配进响应函
        assert "无偏离声明。" not in text  # 商务偏差内容不混入响应函
        inner2 = zipfile.ZipFile(io.BytesIO(zf.read("商务文件/（二）商务偏差表.docx")))
        text2 = _wt(etree.fromstring(inner2.read("word/document.xml")))
        assert "无偏离声明。" in text2
        assert "供货期" in text2  # 表格原文保留


def test_export_requires_deliverable_export_permission(client):
    h, pid = _setup(client)
    _add_deliverable(client, h, pid, 1, {"nodes": [{"id": "n1", "text": "商务"}]})

    engine = create_async_engine(f"sqlite+aiosqlite:///{TEST_DB}")
    factory = async_sessionmaker(engine, expire_on_commit=False)

    async def strip():
        async with factory() as session:
            user = await session.scalar(select(AppUser).where(AppUser.email == "x@test.com"))
            user.permissions = ["file.read"]
            await session.commit()

    asyncio.run(strip())
    engine.sync_engine.dispose()

    r = client.post(f"/api/v1/projects/{pid}/export", json={"formats": ["docx"]}, headers=h)
    assert r.status_code == 403
