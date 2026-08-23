"""成文工具链信息信号测试：seal 身份/校验信号 + 打包全量核对信号（只提示、不拦截——
合规性由主会话+验收/评审子 agent 保证，服务端不设硬性流程代码）。"""

from __future__ import annotations

import asyncio
import io
import time
import zipfile

from docx import Document
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

from app.config import settings

TEST_DB = "./.test_bidvolt.db"


def _setup(client):
    r = client.post(
        "/api/v1/auth/register",
        json={"email": "asm@test.com", "password": "Abc12345", "enterprise_name": "成文测试企业"},
    )
    headers = {"Authorization": f"Bearer {r.json()['access_token']}"}
    pid = client.post("/api/v1/projects", json={"name": "成文项目"}, headers=headers).json()["project_id"]
    return headers, pid


def _docx_bytes(paragraphs: list[str]) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_seal_returns_identity_signals(client, monkeypatch):
    """seal 只给信息信号、不拦截：回执带 req_title/matched_title/was_verified，供 agent 自查。"""
    monkeypatch.setattr(settings, "agent_pipeline_enabled", 1)
    _h, pid = _setup(client)

    from app.services import assembly_service

    engine = create_async_engine("sqlite+aiosqlite:///" + TEST_DB)
    maker = async_sessionmaker(engine, expire_on_commit=False)

    src_text = "（一）响应函及报价汇总表（在投标工具在线填写）\n我方承诺……"
    doc = Document()
    for line in src_text.split("\n"):
        doc.add_paragraph(line)
    base = {
        "doc": doc,
        "sess": None,
        "source_text": src_text,
        "file_id": 1,
        "req_id": 1,
        "title": "（一）响应函及报价汇总表",
        "matched_title": "（一）响应函及报价汇总表（在投标工具在线填写）",
        "verified": False,
        "task_id": 1,
        "project_id": pid,
        "enterprise_id": 1,
        "created": time.time(),
    }
    assembly_service._SLICES["stest1"] = dict(base)

    async def _seal(fname):
        async with maker() as session:
            return await assembly_service.seal_slice(session, "stest1", 1, "价格文件", fname)

    # 未 verify 也允许封存（不拦截），但 was_verified 信号如实=false
    res = asyncio.run(_seal("（二）报价明细表.docx"))
    assert res["artifact_id"] > 0
    assert res["was_verified"] is False
    assert res["req_title"] == "（一）响应函及报价汇总表"
    assert res["matched_title"].startswith("（一）响应函及报价汇总表")

    # verify 后封存：was_verified=true
    assembly_service._SLICES["stest2"] = {
        **base,
        "doc": Document(io.BytesIO(_docx_bytes(["（一）响应函及报价汇总表", "我方承诺……"]))),
        "source_text": "（一）响应函及报价汇总表\n我方承诺……",
        "verified": False,
        "created": time.time(),
    }
    r = assembly_service.verify_slice("stest2", 1)
    assert r["passed"] is True
    async def _seal2():
        async with maker() as session:
            return await assembly_service.seal_slice(session, "stest2", 1, "价格文件", "（一）响应函及报价汇总表.docx")

    res2 = asyncio.run(_seal2())
    assert res2["was_verified"] is True
    assert res2["matched_title"].startswith("（一）响应函及报价汇总表")
    assembly_service._SLICES.pop("stest1", None)
    asyncio.run(engine.dispose())


def _seed_pkg(maker, pid, artifacts):
    from app.models.agent import AgentArtifact
    from app.models.requirement import Requirement
    from app.models.task import Task

    async def _run():
        async with maker() as session:
            task = Task(
                enterprise_id=1,
                project_id=pid,
                task_type="agent_pipeline",
                idempotency_key=f"asm-pkg-{len(artifacts)}-{time.time_ns()}",
                status=2,
                payload={},
            )
            session.add(task)
            await session.flush()
            for i, (title, role, order) in enumerate(
                [
                    ("（一）响应函及报价汇总表", "price", 1),
                    ("（二）报价明细表", "price", 2),
                    ("一、价格文件", "price", 0),
                ]
            ):
                session.add(
                    Requirement(
                        enterprise_id=1,
                        project_id=pid,
                        req_type="doc_template",
                        content=title,
                        current=True,
                        structured={"role": role, "order": order},
                    )
                )
            for name, content in artifacts:
                session.add(
                    AgentArtifact(
                        enterprise_id=1,
                        project_id=pid,
                        task_id=task.id,
                        kind="item_docx",
                        name=name,
                        mime="application/octet-stream",
                        content=content,
                    )
                )
            await session.commit()
            return task.id

    return asyncio.run(_run())


def _pack(maker, pid, task_id):
    from app.services import assembly_service

    async def _run():
        async with maker() as session:
            from sqlalchemy import text

            ids = [
                r[0]
                for r in (
                    await session.execute(
                        text("select id from agent_artifact where task_id=:t and kind='item_docx'"),
                        {"t": task_id},
                    )
                ).fetchall()
            ]
            return await assembly_service.package_zip(session, 1, pid, task_id, ids, None)

    return asyncio.run(_run())


def test_package_zip_reports_missing_item(client, monkeypatch):
    """打包只给信号、不拦截：缺条目照常出包，回执如实返回 missing_file_items。"""
    monkeypatch.setattr(settings, "agent_pipeline_enabled", 1)
    _h, pid = _setup(client)
    engine = create_async_engine("sqlite+aiosqlite:///" + TEST_DB)
    maker = async_sessionmaker(engine, expire_on_commit=False)
    tid = _seed_pkg(
        maker, pid,
        [("价格文件/（一）响应函及报价汇总表.docx",
          _docx_bytes(["（一）响应函及报价汇总表", "我方承诺……"]))],
    )
    result = _pack(maker, pid, tid)
    assert "（二）报价明细表" in result["missing_file_items"]
    assert result["audit"]["coverage_ok"] is False
    asyncio.run(engine.dispose())


def test_package_zip_reports_duplicate_and_identity(client, monkeypatch):
    """打包只给信号、不拦截：同部分雷同与身份不符如实进入 audit 信号。"""
    monkeypatch.setattr(settings, "agent_pipeline_enabled", 1)
    _h, pid = _setup(client)
    engine = create_async_engine("sqlite+aiosqlite:///" + TEST_DB)
    maker = async_sessionmaker(engine, expire_on_commit=False)

    # 雷同：两份不同文件名、同一内容
    tid = _seed_pkg(
        maker, pid,
        [
            ("价格文件/（一）响应函及报价汇总表.docx", _docx_bytes(["（一）响应函及报价汇总表", "我方承诺……"])),
            ("价格文件/（二）报价明细表.docx", _docx_bytes(["（一）响应函及报价汇总表", "我方承诺……"])),
        ],
    )
    result = _pack(maker, pid, tid)
    assert result["audit"]["unique_ok"] is False
    assert result["audit"]["duplicate_pairs"], result["audit"]

    # 身份不符：文件名是报价明细表，内容却是响应函
    tid2 = _seed_pkg(
        maker, pid,
        [
            ("价格文件/（一）响应函及报价汇总表.docx", _docx_bytes(["（一）响应函及报价汇总表", "我方承诺……"])),
            ("价格文件/（二）报价明细表.docx", _docx_bytes(["响应函价格表", "合计报价……"])),
        ],
    )
    result2 = _pack(maker, pid, tid2)
    assert result2["audit"]["identity_ok"] is False
    assert result2["audit"]["identity_issues"], result2["audit"]
    asyncio.run(engine.dispose())


def test_package_zip_passes_full_audit_and_dedupes(client, monkeypatch):
    """打包核对信号全绿：产出 zip，同名产物改名不覆盖，manifest 带 audit 结论。"""
    monkeypatch.setattr(settings, "agent_pipeline_enabled", 1)
    _h, pid = _setup(client)
    engine = create_async_engine("sqlite+aiosqlite:///" + TEST_DB)
    maker = async_sessionmaker(engine, expire_on_commit=False)
    # 两份正确内容 + 一份与首份同名但正文不同的多余产物（测改名去重：审计通过后同名改名）
    tid = _seed_pkg(
        maker, pid,
        [
            ("价格文件/（一）响应函及报价汇总表.docx", _docx_bytes(["（一）响应函及报价汇总表", "我方承诺……"])),
            ("价格文件/（二）报价明细表.docx", _docx_bytes(["（二）报价明细表", "明细报价如下……"])),
            ("价格文件/（一）响应函及报价汇总表.docx", _docx_bytes(["（一）响应函及报价汇总表", "我方承诺（补充）……"])),
        ],
    )
    result = _pack(maker, pid, tid)
    assert result.get("missing_file_items") == []

    from sqlalchemy import text

    async def _load():
        async with maker() as session:
            row = (
                await session.execute(
                    text("select content from agent_artifact where kind='zip' and task_id=:t"),
                    {"t": tid},
                )
            ).fetchone()
            return row[0]

    zdata = asyncio.run(_load())
    with zipfile.ZipFile(io.BytesIO(zdata)) as zf:
        names = zf.namelist()
        assert names.count("价格文件/（一）响应函及报价汇总表.docx") == 1
        assert any("(2)" in n for n in names), names
        assert "价格文件/（二）报价明细表.docx" in names
        assert "会话记录/主会话记录.md" in names
        import json

        manifest = json.loads(zf.read("manifest.json"))
        assert manifest["audit"]["coverage_ok"] is True
        assert manifest["audit"]["unique_ok"] is True
        assert manifest["audit"]["identity_ok"] is True
    asyncio.run(engine.dispose())


def test_inspect_artifact_previews(client, monkeypatch):
    """产物自检：xlsx 预览行列与行内容；zip 预览文件清单。"""
    monkeypatch.setattr(settings, "agent_pipeline_enabled", 1)
    h, pid = _setup(client)

    from app.models.agent import AgentArtifact
    from app.models.task import Task
    from app.services import assembly_service

    engine = create_async_engine("sqlite+aiosqlite:///" + TEST_DB)
    maker = async_sessionmaker(engine, expire_on_commit=False)

    async def _seed():
        async with maker() as session:
            task = Task(
                enterprise_id=1,
                project_id=pid,
                task_type="agent_pipeline",
                idempotency_key="asm-task-2",
                status=2,
                payload={},
            )
            session.add(task)
            await session.flush()
            # 一个真实 xlsx 产物（多行多列）
            from openpyxl import Workbook

            wb = Workbook()
            ws = wb.active
            ws.title = "报价单"
            ws.append(["序号", "名称"])
            ws.append([1, "虚拟电厂平台"])
            buf = io.BytesIO()
            wb.save(buf)
            session.add(
                AgentArtifact(
                    enterprise_id=1,
                    project_id=pid,
                    task_id=task.id,
                    kind="xlsx",
                    name="价格文件/报价单.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    content=buf.getvalue(),
                )
            )
            await session.commit()
            return task.id

    task_id = asyncio.run(_seed())

    async def _inspect():
        async with maker() as session:
            from sqlalchemy import text

            aid = (
                await session.execute(
                    text("select id from agent_artifact where task_id=:t and kind='xlsx'"),
                    {"t": task_id},
                )
            ).fetchone()[0]
            return await assembly_service.inspect_artifact(session, 1, pid, task_id, aid)

    info = asyncio.run(_inspect())
    print("inspect:", info)
    assert info["kind"] == "xlsx"
    assert info["sheets"][0]["rows"] == 2
    assert info["sheets"][0]["cols"] == 2
    assert "虚拟电厂平台" in str(info["sheets"][0]["preview"])
    asyncio.run(engine.dispose())
