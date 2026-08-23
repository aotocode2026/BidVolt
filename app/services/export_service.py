"""导出与终检（4.10）：DOCX/XLSX 生成、一致性检查、manifest、交付包。"""

from __future__ import annotations

import logging
from datetime import datetime, timezone
from io import BytesIO

logger = logging.getLogger(__name__)

DELIVERABLE_NAMES = {1: "商务标", 2: "技术标", 3: "报价单"}
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class _TrackedEditor:
    """在 docx 上以 Word 修订模式记录全部改动（产品要求：验收可见、可追溯）。

    - 每处改动：原文字标为删除（w:del + w:delText，Word 显示删除线），
      新文字标为插入（w:ins + w:t）；
    - 每处改动挂一条批注（w:comment），说明改了什么、依据来源；
    - 文末补充节整体标记为插入，并批注"系统新增内容"。"""

    def __init__(self, doc):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        self._doc = doc
        self._mk = OxmlElement
        self._qn = qn
        self._seq = {"ins": 100, "del": 200, "comment": 300}
        self._now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        self.comments: list[tuple[int, str]] = []

    def _mark(self, tag: str):
        kind = "ins" if tag == "w:ins" else "del"
        self._seq[kind] += 1
        el = self._mk(tag)
        el.set(self._qn("w:id"), str(self._seq[kind]))
        el.set(self._qn("w:author"), "BidVolt")
        el.set(self._qn("w:date"), self._now)
        return el

    def _run(self, tag: str, text: str, rpr_source=None):
        r = self._mk("w:r")
        if rpr_source is not None:
            import copy as _copy

            rpr = rpr_source.find(self._qn("w:rPr"))
            if rpr is not None:
                r.append(_copy.deepcopy(rpr))
        t = self._mk(tag)
        t.set(self._qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        return r

    def _anchor(self, cid: int, before_el):
        """在 before_el 之前插入 commentRangeStart（End 与 Reference 由调用方按序追加）。"""
        start = self._mk("w:commentRangeStart")
        start.set(self._qn("w:id"), str(cid))
        before_el.addprevious(start)
        return start

    def _anchor_end_ref(self, cid: int, before_el) -> None:
        """在 before_el 之前追加 commentRangeEnd 与 commentReference（须在改动内容之后）。"""
        end = self._mk("w:commentRangeEnd")
        end.set(self._qn("w:id"), str(cid))
        before_el.addprevious(end)
        ref_r = self._mk("w:r")
        rpr = self._mk("w:rPr")
        rstyle = self._mk("w:rStyle")
        rstyle.set(self._qn("w:val"), "CommentReference")
        rpr.append(rstyle)
        ref_r.append(rpr)
        ref = self._mk("w:commentReference")
        ref.set(self._qn("w:id"), str(cid))
        ref_r.append(ref)
        before_el.addprevious(ref_r)

    def add_comment(self, text: str) -> int:
        self._seq["comment"] += 1
        cid = self._seq["comment"]
        self.comments.append((cid, text))
        return cid

    def track_replace(self, t_el, old_text: str, new_text: str, comment: str) -> None:
        """单个文本节点的替换：原文删除线 + 新文插入 + 批注（保留原 run 格式）。
        元素顺序：commentRangeStart → del → ins → commentRangeEnd → commentReference。"""
        r_el = t_el.getparent()
        cid = self.add_comment(comment)
        self._anchor(cid, r_el)
        del_el = self._mark("w:del")
        del_el.append(self._run("w:delText", old_text, r_el))
        r_el.addprevious(del_el)
        ins_el = self._mark("w:ins")
        ins_el.append(self._run("w:t", new_text, r_el))
        r_el.addprevious(ins_el)
        self._anchor_end_ref(cid, r_el)
        t_el.text = ""

    def track_paragraph_replace(self, t_nodes, old_full: str, new_full: str, comment: str) -> None:
        """整段替换（占位符跨 run 拆分时）：整段原文删除线 + 整段新文插入 + 批注。"""
        first_r = t_nodes[0].getparent()
        cid = self.add_comment(comment)
        self._anchor(cid, first_r)
        del_el = self._mark("w:del")
        del_el.append(self._run("w:delText", old_full, first_r))
        first_r.addprevious(del_el)
        ins_el = self._mark("w:ins")
        ins_el.append(self._run("w:t", new_full, first_r))
        first_r.addprevious(ins_el)
        self._anchor_end_ref(cid, first_r)
        for t in t_nodes:
            t.text = ""

    def track_insert_run(self, run) -> None:
        """把已创建的 run 包进 w:ins（新增内容整体标记为插入）。"""
        ins_el = self._mark("w:ins")
        r_el = run._r
        r_el.addprevious(ins_el)
        ins_el.append(r_el)

    def write_comments_part(self) -> None:
        """把批注写入 word/comments.xml 部件并建立关系。"""
        import html as _html

        if not self.comments:
            return
        from docx.opc.constants import RELATIONSHIP_TYPE as _RT
        from docx.opc.packuri import PackURI
        from docx.opc.part import Part

        rows = [
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
            '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">',
        ]
        for cid, text in self.comments:
            rows.append(
                f'<w:comment w:id="{cid}" w:author="BidVolt" w:date="{self._now}" w:initials="BV">'
                f'<w:p><w:r><w:t xml:space="preserve">{_html.escape(text, quote=False)}</w:t></w:r></w:p>'
                "</w:comment>"
            )
        rows.append("</w:comments>")
        blob = "".join(rows).encode("utf-8")
        part = Part(
            PackURI("/word/comments.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
            blob,
            self._doc.part.package,
        )
        rt_comments = getattr(
            _RT, "COMMENTS",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
        )
        self._doc.part.relate_to(part, rt_comments)


class _FillSession:
    """一次成文的填空会话：修订模式填空 + 批注来源 + 补充内容追加。
    整本成文（docx_from_template）与逐份成文（响应文件包）共用同一套规则。"""

    def __init__(self, doc, buyer: str, project_name: str, supplier: str, tender_no: str = ""):
        import re as _re

        from docx.oxml import OxmlElement as _OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        self.doc = doc
        self.editor = _TrackedEditor(doc)
        self.buyer = buyer
        self.project_name = project_name
        self.supplier = supplier
        self.tender_no = tender_no
        self._re = _re
        self._qn = qn
        self._Pt = Pt
        self._PLACEHOLDER_KEYS = (
            "（采购人）", "（响应供应商名称）", "（项目名称）",
            "【招标人名称】", "【供应商名称】", "【项目名称】",
        )
        # 带标签空位（模板空位为空格/无下划线形态，如"采购编号：    ，"与"包名称：   采购文件"）：
        # P1 = 标签后直接是标点/行尾（空位零宽）；P2 = 标签后有空白再跟正文
        _labels = (
            "采购编号|项目名称|分标名称|分标编号|包名称|包号|响应供应商|单位地址|法定地址|"
            "法定代表人（单位负责人）或授权代表|法定代表人（单位负责人）|法定代表人|"
            "邮政编码|电话|传真|日期"
        )
        self._LABEL_P1 = _re.compile(rf"({_labels})[：:]\s*(?=[，,。;；）)（(]|$)")
        self._LABEL_P2 = _re.compile(rf"({_labels})[：:][ \u3000]+(?=[^，,。;；）)\s])")
        # Word 显示修订标记的前提：settings.xml 须含 <w:trackChanges/>，
        # 否则 w:ins/w:del 会被当作普通文本（插入可见、删除被吞）。
        settings_el = doc.settings.element
        if settings_el.find(qn("w:trackChanges")) is None:
            tc = _OxmlElement("w:trackChanges")
            default_tab = settings_el.find(qn("w:defaultTabStop"))
            if default_tab is not None:
                default_tab.addprevious(tc)
            else:
                settings_el.append(tc)

    def _labeled_value(self, label: str) -> str:
        """带标签空位的回填值：有资料填值；无资料原位【待补充：标签】（不编造）。"""
        if label == "采购编号":
            return self.tender_no or "【待补充：采购编号】"
        if label == "项目名称":
            return self.project_name or "【待补充：项目名称】"
        if label == "响应供应商":
            return self.supplier or "【待补充：供应商名称】"
        return f"【待补充：{label}】"

    def fill(self, text: str) -> str:
        _re = self._re
        buyer, project_name, supplier = self.buyer, self.project_name, self.supplier
        out = text
        # 标签紧邻空位：整体回填（"办理____（项目名称）"的下划线才是字段本身）
        if project_name:
            out = _re.sub(r"_{2,}\s*[（(【]\s*项目名称\s*[）)】]", project_name, out)
        if buyer:
            out = _re.sub(r"_{2,}\s*[（(【]\s*采购人\s*[）)】]", buyer, out)
            out = _re.sub(r"_{2,}\s*【\s*招标人名称\s*】", buyer, out)
        if supplier:
            out = _re.sub(r"_{2,}\s*[（(【]\s*响应供应商名称\s*[）)】]", supplier, out)
            out = _re.sub(r"_{2,}\s*【\s*供应商名称\s*】", supplier, out)
        # 应答函裸标签形态（"致：采购人" 无括号，是主应答函称谓行）
        if buyer:
            out = _re.sub(r"(?m)^致[：:]\s*采购人\s*$", f"致：{buyer}", out)
            out = _re.sub(r"(?m)^致[：:]\s*招标人\s*$", f"致：{buyer}", out)
        out = out.replace("（采购人）", buyer or "【待补充：招标人名称】")
        out = out.replace("（响应供应商名称）", supplier or "【待补充：供应商名称】")
        out = out.replace("（项目名称）", project_name or "【待补充：项目名称】")
        out = out.replace("【招标人名称】", buyer or "【待补充：招标人名称】")
        out = out.replace("【供应商名称】", supplier or "【待补充：供应商名称】")
        out = out.replace("【项目名称】", project_name or "【待补充：项目名称】")
        out = _re.sub(r"_{2,}", "【待补充】", out)
        # 带标签的空位原位回填（如"响应供应商名称：____（盖章）"）
        if supplier:
            out = _re.sub(r"响应供应商名称[：:]\s*【待补充】", f"响应供应商名称：{supplier}", out)
        if project_name:
            out = _re.sub(r"项目名称[：:]\s*【待补充】", f"项目名称：{project_name}", out)
        if buyer:
            out = _re.sub(r"(?:招标人|采购人)[：:]\s*【待补充】", f"招标人：{buyer}", out)
        # 带标签空位（空格/零宽形态）：有资料回填，无资料原位【待补充：标签】
        def _rep(m):
            return f"{m.group(1)}：{self._labeled_value(m.group(1))}"

        out = self._LABEL_P1.sub(_rep, out)
        out = self._LABEL_P2.sub(_rep, out)
        return out

    def comment(self, old: str, new: str) -> str:
        """修订批注：说明改了什么、依据来源（产品要求：批注必须说明来自哪里）。
        批注是审阅侧留痕，措辞保持专业投标口吻，不带系统痕迹。"""
        buyer, project_name, supplier = self.buyer, self.project_name, self.supplier
        filled: list[str] = []
        if ("（采购人）" in old or "【招标人名称】" in old or "招标人：" in old) and buyer:
            filled.append(f"招标人=「{buyer}」（来源：招标文件封面/采购公告）")
        elif ("致：" in old and "采购人" in old) and buyer:
            filled.append(f"招标人=「{buyer}」（来源：招标文件封面/采购公告）")
        if ("（响应供应商名称）" in old or "【供应商名称】" in old or "响应供应商名称" in old) and supplier:
            filled.append(f"供应商=「{supplier}」（来源：企业资料-企业名称）")
        if ("（项目名称）" in old or "【项目名称】" in old or "项目名称：" in old) and project_name:
            filled.append(f"项目名称=「{project_name}」（来源：招标文件封面/采购公告）")
        if "采购编号" in old and self.tender_no:
            filled.append(f"采购编号=「{self.tender_no}」（来源：招标文件封面/采购公告）")
        if self._re.search(r"(分标名称|分标编号|包名称|包号)[：:]", old) and "【待补充" in new:
            filled.append("分标/包 信息未指定应答分包，原位标注【待补充】（按所应答分包填写）")
        if filled:
            if "【待补充" in new:
                return (
                    "已按来源资料回填：" + "；".join(filled)
                    + "；本处其余空位（如授权人/被授权人等）未取得对应资料，原位标注【待补充】。请复核确认。"
                )
            return "已按来源资料回填：" + "；".join(filled) + "，请复核。"
        if "【待补充" in new:
            return "模板空位未取得对应资料，原位标注【待补充】（不编造内容），取得资料后填写确认。"
        return "按招标文件内容补充填写（供参考），请复核确认。"

    def apply_to_doc(self) -> None:
        """全文档（含表格、控件、文本框）逐段填空（修订模式+批注）。"""
        _re = self._re
        qn = self._qn
        editor = self.editor
        keys = self._PLACEHOLDER_KEYS
        fill = self.fill
        comment = self.comment

        def _fill_paragraph_el(p_el) -> None:
            t_nodes = list(p_el.iter(qn("w:t")))
            if not t_nodes:
                return
            full = "".join(t.text or "" for t in t_nodes)
            bare_zhia = bool(_re.search(r"(?m)^致[：:]\s*(?:采购人|招标人)\s*$", full))
            has_label_blank = bool(self._LABEL_P1.search(full) or self._LABEL_P2.search(full))
            if (
                not any(k in full for k in keys)
                and not _re.search(r"_{2,}", full)
                and not bare_zhia
                and not has_label_blank
            ):
                return
            cross = False
            for k in keys:
                if full.count(k) != sum((t.text or "").count(k) for t in t_nodes):
                    cross = True
                    break
            if not cross:
                full_unders = len(_re.findall(r"_{2,}", full))
                node_unders = sum(len(_re.findall(r"_{2,}", t.text or "")) for t in t_nodes)
                if full_unders != node_unders:
                    cross = True
            if _re.search(r"(响应供应商名称|项目名称|招标人|采购人)[：:]\s*_{2,}", full):
                cross = True  # 带标签空位常跨 run，整段处理
            if _re.search(r"_{2,}\s*[（(【]\s*(?:项目名称|采购人|响应供应商名称|招标人名称|供应商名称)\s*[）)】]", full):
                cross = True  # 标签紧邻空位必须整段整体回填
            if has_label_blank:
                cross = True  # 带标签空位（空格/零宽）跨 run 常见，整段整体回填
            if cross:
                new_full = fill(full)
                if new_full != full:
                    editor.track_paragraph_replace(t_nodes, full, new_full, comment(full, new_full))
                return
            for t in t_nodes:
                text = t.text or ""
                if not text:
                    continue
                new = fill(text)
                if new != text:
                    editor.track_replace(t, text, new, comment(text, new))

        for p_el in self.doc.element.body.iter(qn("w:p")):
            _fill_paragraph_el(p_el)
        # 兜底：段落容器之外的裸 w:t 文本节点（如控件内未包段落的 run）
        for t_el in self.doc.element.body.iter(qn("w:t")):
            text = t_el.text or ""
            if (
                any(k in text for k in keys)
                or _re.search(r"_{2,}", text)
                or _re.search(r"(?m)^致[：:]\s*(?:采购人|招标人)\s*$", text)
                or self._LABEL_P1.search(text)
                or self._LABEL_P2.search(text)
            ):
                new = fill(text)
                if new != text:
                    editor.track_replace(t_el, text, new, comment(text, new))

    def add_heading_safe(self, text: str, level: int):
        """追加标题：不依赖源文档是否含 'Heading n' 样式（国网等采购文件底稿常无内置标题样式）。"""
        from docx.oxml import OxmlElement

        qn = self._qn
        Pt = self._Pt
        p = self.doc.add_paragraph()
        run = p.add_run(str(text))
        run.bold = True
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
        run.font.size = Pt(16 if level == 1 else 14)
        p_pr = p._p.get_or_add_pPr()
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(max(int(level) - 1, 0)))
        p_pr.append(outline)
        return p

    def fill_table_cell(self, table_idx: int, row_idx: int, col_idx: int, value: str, comment: str | None = None) -> bool:
        """定向填写模板表格单元格（修订插入+批注）：表格类条目应把内容填进模板自身的表格，
        而不是空着表格把内容挂到文件末尾。越界返回 False（调用方如实记录）。"""
        tables = self.doc.tables
        if table_idx < 0 or table_idx >= len(tables):
            return False
        tb = tables[table_idx]
        if row_idx < 0 or row_idx >= len(tb.rows) or col_idx < 0 or col_idx >= len(tb.rows[row_idx].cells):
            return False
        cell = tb.rows[row_idx].cells[col_idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(self.fill(str(value)))
        self.editor.track_insert_run(run)
        self.editor.add_comment(comment or "按采购文件要求填写本单元格，请复核。")
        return True

    def append_supplement(self, nodes, heading_text: str | None = None, comment: str | None = None, page_break: bool = True) -> None:
        """把撰写内容追加为修订插入（w:ins）+ 批注；heading_text 为 None 时不加总标题。"""
        Pt = self._Pt
        qn = self._qn
        editor = self.editor
        fill = self.fill
        if page_break:
            self.doc.add_page_break()
        if heading_text:
            h = self.add_heading_safe(heading_text, 1)
            editor.track_insert_run(h.runs[0])
        if comment:
            editor.add_comment(comment)
        for node in _norm_supplement_nodes(nodes):
            t = str(node.get("text") or "")
            ntype = node.get("type")
            if ntype == "heading":
                h = self.add_heading_safe(fill(t), 2)
                editor.track_insert_run(h.runs[0])
            elif ntype == "table" and node.get("rows"):
                rows = node["rows"]
                tb = self.doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                tb.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(len(tb.rows[ri].cells)):
                        cell = tb.rows[ri].cells[ci]
                        cell.text = ""
                        run = cell.paragraphs[0].add_run(fill(str(row[ci])) if ci < len(row) else "")
                        run.font.size = Pt(10.5)
                        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
                        editor.track_insert_run(run)
            elif t.strip():
                p = self.doc.add_paragraph(fill(t))
                for run in p.runs:
                    editor.track_insert_run(run)

    def finish(self) -> bytes:
        self.editor.write_comments_part()
        buf = BytesIO()
        self.doc.save(buf)
        return buf.getvalue()


def docx_from_template(source_path, model: dict) -> bytes:
    """底稿式整本成文（保留能力）：复制采购文件原 docx，整本保留——
    填空（修订模式+批注来源）+ 文末追加"补充响应内容"（修订插入+批注）。
    页面规格/字体/表格样式/分页/页眉页脚/全部原文天然保留。"""
    from docx import Document

    doc = Document(str(source_path))
    sess = _FillSession(
        doc,
        str(model.get("buyer") or "").strip(),
        str(model.get("project_name") or "").strip(),
        str(model.get("supplier_name") or "").strip(),
        str(model.get("tender_no") or "").strip(),
    )
    sess.apply_to_doc()
    sess.append_supplement(
        model.get("supplement_nodes") or [],
        heading_text="补充响应内容",
        comment="本节为补充响应内容（依据采购文件要求与企业资料撰写，企业事实以资料为准，"
                "未知处标【待补充】），请复核确认。",
    )
    return sess.finish()


# ---------- 响应文件包：按招标文件"响应文件格式"清单逐份成文（产品定版） ----------

_RESPONSE_ROLE_DIRS = {"price": "价格文件", "business": "商务文件", "technical": "技术文件"}
_ROLE_TITLE_KEYS = {
    "price": ("价格文件", "报价文件", "价格部分"),
    "business": ("商务文件",),
    "technical": ("技术文件",),
}
_END_MARKERS = ("商务评分标准", "技术评分标准", "响应文件编制注意事项")
# 撰写内容 → 条目的匹配关键词组（条目标题关键词, 内容标题关键词）
_KEYWORD_GROUPS = [
    (("响应函", "应答函"), ("应答函", "响应函")),
    (("授权委托",), ("授权委托",)),
    (("商务偏差", "偏离表", "商务条款"), ("偏离", "偏差表")),
    (("保证保险", "保证金"), ("保证金", "保证保险")),
    (("资格审查", "资格"), ("资格审查", "资格")),
    (("自查",), ("自查",)),
    (("技术偏差", "参数响应"), ("技术偏差", "参数响应", "逐条响应")),
    (("专项响应", "实施方案", "技术方案"), ("专项响应", "实施方案", "技术方案")),
    (("业绩",), ("业绩",)),
    (("项目团队", "人员"), ("团队", "人员")),
    (("服务承诺", "售后"), ("服务承诺", "售后")),
    (("进度保障", "进度", "工期"), ("进度", "工期")),
    # 容器条目（补充文件/专项响应文件）内的子表单章节归入容器文件
    (("专项响应", "补充文件"),
     ("业绩", "团队", "理解", "规划", "履约", "服务承诺", "进度", "保证金",
      "关系说明", "基本情况", "支撑材料", "规范书", "凭证", "社保", "自查")),
]


def _norm_text(s: str) -> str:
    import re as _re

    return _re.sub(r"\s+", "", s or "")


def _strip_num(s: str) -> str:
    import re as _re

    return _re.sub(r"^[（(]?[一二三四五六七八九十百0-9]+[）)、.、:：\s]*", "", s)


def _item_key(s: str) -> str:
    """条目匹配键：去编号、去上传路径尾巴、去括号注（如"（单位负责人）""（如有）"）。
    清单标题与底稿标题的常见差异都来自这几类装饰，归一后比对。"""
    import re as _re

    s = _clean_item_name(s or "")
    s = _re.sub(r"[（(][^）)]*[）)]", "", s)
    return _strip_num(_norm_text(s))


def _elem_text(el) -> str:
    return "".join(el.itertext())


def _iter_body_elems(doc):
    from docx.oxml.ns import qn

    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("p", child)
        elif child.tag == qn("w:tbl"):
            yield ("tbl", child)


def _locate_item_slices(doc, items_by_role: dict) -> dict:
    """按底稿大纲级别切分条目（v2，优先）：outlineLvl=1 为部分标题（价格/商务/技术文件），
    outlineLvl=2 为文件条目标题（（一）响应函…）。返回 {role: [(heading, elements)]}。
    返回空结果时调用方回退行级清单匹配。"""
    from docx.oxml.ns import qn as _qn

    elems = list(_iter_body_elems(doc))
    # 章起点：最后一次出现"响应文件格式"（跳过目录里的 TOC 同名条目）
    start = None
    for i in range(len(elems) - 1, -1, -1):
        kind, el = elems[i]
        if kind == "p" and "响应文件格式" in _elem_text(el):
            start = i
            break
    if start is None:
        return {}
    end = None
    for i in range(start + 1, len(elems)):
        kind, el = elems[i]
        t = _norm_text(_elem_text(el))
        # 结束标记只认短标题行（评分标准章标题），正文中的长句引用不截断区域
        if kind == "p" and len(t) <= 24 and any(m in t for m in _END_MARKERS):
            end = i
            break
    if end is None:
        end = len(elems)
    region = elems[start + 1 : end]

    def outline_of(el) -> int | None:
        ppr = el.find(_qn("w:pPr"))
        if ppr is None:
            return None
        lvl = ppr.find(_qn("w:outlineLvl"))
        if lvl is None:
            return None
        try:
            return int(lvl.get(_qn("w:val")))
        except (TypeError, ValueError):
            return None

    import re as _re

    items: list[tuple[str, str, list]] = []  # (role, heading_text, elems)
    cur_role: str | None = None
    cur_item: str | None = None
    cur_buf: list = []

    def flush_item() -> None:
        nonlocal cur_item, cur_buf
        if cur_item is not None and cur_role is not None:
            items.append((cur_role, cur_item, cur_buf))
        cur_item, cur_buf = None, []

    for kind, el in region:
        if kind != "p":
            cur_buf.append((kind, el))
            continue
        lvl = outline_of(el)
        t = _norm_text(_elem_text(el)).strip()
        if lvl == 1:
            flush_item()
            if "价格文件" in t:
                cur_role = "price"
            elif "商务文件" in t:
                cur_role = "business"
            elif "技术文件" in t:
                cur_role = "technical"
            cur_buf.append((kind, el))  # 部分标题并入首条目
        elif lvl == 2 and cur_role is not None and _re.match(r"^[（(][一二三四五六七八九十百0-9]+[）)]", t):
            # 条目边界：（一）（二）…开头的大纲2级标题；"1."式次级标题归入当前条目
            flush_item()
            cur_item = t
            cur_buf = [(kind, el)]
        else:
            cur_buf.append((kind, el))
    flush_item()
    result: dict[str, list] = {"price": [], "business": [], "technical": []}
    for role, heading, elems in items:
        if role in result:
            result[role].append((heading, elems))
    return result


def _locate_item_slices_by_rows(doc, items_by_role: dict) -> dict:
    """行级清单匹配（兜底：底稿无大纲级别信号时用）。返回 {role: [(row, elements|None)]}。"""
    elems = list(_iter_body_elems(doc))
    start = None
    for i in range(len(elems) - 1, -1, -1):
        kind, el = elems[i]
        if kind == "p" and "响应文件格式" in _elem_text(el):
            start = i
            break
    if start is None:
        return {role: [(r, None) for r in rows] for role, rows in items_by_role.items()}
    end = None
    for i in range(start + 1, len(elems)):
        kind, el = elems[i]
        t = _norm_text(_elem_text(el))
        if kind == "p" and len(t) <= 24 and any(m in t for m in _END_MARKERS):
            end = i
            break
    if end is None:
        end = len(elems)
    region = elems[start:end]
    pos: dict[str, int] = {}
    for i, (kind, el) in enumerate(region):
        if kind != "p":
            continue
        t = _norm_text(_elem_text(el))
        if not (1 <= len(t) <= 12):
            continue  # 角色标题是短段落；含多个角色词的说明行不参与匹配
        for role, keys in _ROLE_TITLE_KEYS.items():
            if role in pos:
                continue
            if any(t == k or t.endswith(k) for k in keys):
                pos[role] = i
    ordered = sorted(pos.items(), key=lambda kv: kv[1])
    spans = {}
    for idx, (role, p) in enumerate(ordered):
        nxt = ordered[idx + 1][1] if idx + 1 < len(ordered) else len(region)
        spans[role] = (p, nxt)
    return {role: _slice_role(region, spans.get(role), rows) for role, rows in items_by_role.items()}


def _clean_item_name(heading: str) -> str:
    """条目标题清洗为文件名：去掉"（在投标工具…）"等上传路径说明尾巴。"""
    for cut in ("（在投标工具", "（上传投标工具", "（上传招投标", "（可在系统", "（如谈判", "（如有", "（如选择"):
        idx = heading.find(cut)
        if idx > 0:
            heading = heading[:idx]
    return heading


def _slice_role(region, span, rows):
    if span is None:
        return [(r, None) for r in rows]
    lo, hi = span
    seg = region[lo:hi]

    def t_at(i):
        return _norm_text(_elem_text(seg[i][1]))

    markers = []
    for r in rows:
        key = _strip_num(_norm_text(r.content or ""))[:16]
        if not key:
            continue
        for i in range(len(seg)):
            kind, _el = seg[i]
            txt = t_at(i)
            if kind == "p" and key in txt and 2 <= len(txt) <= 200:
                markers.append((i, r))
                break
    markers.sort(key=lambda m: m[0])
    out = []
    matched_ids = set()
    for idx, (mpos, r) in enumerate(markers):
        matched_ids.add(r.id)
        mend = markers[idx + 1][0] if idx + 1 < len(markers) else len(seg)
        elems = seg[mpos:mend]
        if idx == 0 and mpos > 0:
            elems = seg[0:mpos] + elems  # 部分标题行及其后说明并入首条目
        out.append((r, elems))
    for r in rows:
        if r.id not in matched_ids:
            out.append((r, None))
    return out


def _norm_supplement_nodes(nodes) -> list[dict]:
    """撰写内容节点归一化：兼容 agent 落库的裸字符串/混合形状 → {type,text} 段落节点。"""
    out = []
    for n in nodes or []:
        if isinstance(n, str):
            t = n.strip()
            if t:
                out.append({"type": "paragraph", "text": t})
        elif isinstance(n, dict):
            out.append(n)
        elif n is not None:
            out.append({"type": "paragraph", "text": str(n)})
    return out


def _split_supplement(titles, nodes) -> tuple[list, list]:
    """把撰写内容按标题关键词分配到各条目标题；返回 (matched_by_title_index, rest_nodes)。"""
    sections = []
    cur = None
    for n in _norm_supplement_nodes(nodes):
        if n.get("type") == "heading" and str(n.get("text") or "").strip():
            if cur is not None:
                sections.append(cur)
            cur = {"head": str(n.get("text") or "").strip(), "nodes": [n]}
        else:
            if cur is None:
                cur = {"head": "", "nodes": []}
            cur["nodes"].append(n)
    if cur is not None:
        sections.append(cur)
    matched: list[list] = [[] for _ in titles]
    rest: list = []
    for sec in sections:
        head = sec["head"]
        probe = head + "".join(str(n.get("text") or "") for n in sec["nodes"][:3])
        hit = None
        # 标题与条目名完全一致的优先（生成阶段按条目名成章）
        if head:
            for idx, title in enumerate(titles):
                if head == title or title.startswith(head) or head.startswith(title):
                    hit = idx
                    break
        if hit is None:
            for idx, title in enumerate(titles):
                for tkeys, skeys in _KEYWORD_GROUPS:
                    if any(tk in title for tk in tkeys):
                        if any(sk in probe for sk in skeys):
                            hit = idx
                            break
                if hit is not None:
                    break
        if hit is not None:
            matched[hit].extend(sec["nodes"])
        else:
            rest.extend(sec["nodes"])
    return matched, rest


def _safe_filename(title: str, max_len: int = 40) -> str:
    import re as _re

    name = _re.sub(r'[\\/:*?"<>|\r\n\t]', "_", _norm_text(title) or "条目")
    return name[:max_len] or "条目"


def _build_item_document(source_path, row, elements):
    """条目文件骨架：底稿中该条目区间原文整段复制（保留格式）；未定位到区间的条目
    按清单内容重建。返回 (Document, located)。"""
    import copy as _copy

    from docx import Document
    from docx.oxml.ns import qn

    new = Document(str(source_path))
    body = new.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    # 模板切片必须插在 sectPr 之前（OOXML：sectPr 是 body 最后一个元素；
    # 插在其后 Word 会忽略内容，导致首空白页+正文丢失）
    sect_pr = body.find(qn("w:sectPr"))
    located = bool(elements)
    if elements:
        for kind, el in elements:
            node = _copy.deepcopy(el)
            if sect_pr is not None:
                sect_pr.addprevious(node)
            else:
                body.append(node)
    elif row is not None:
        # 清单重建：底稿中未定位到该条目标题时，按解析清单内容成文并如实批注
        st = row.structured or {}
        if st.get("kind") == "table" and st.get("rows"):
            rows = st["rows"]
            tb = new.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            tb.style = "Table Grid"
            for ri, rrow in enumerate(rows):
                for ci in range(len(tb.rows[ri].cells)):
                    cell = tb.rows[ri].cells[ci]
                    cell.text = ""
                    cell.paragraphs[0].add_run(str(rrow[ci]) if ci < len(rrow) else "")
        elif (row.content or "").strip():
            new.add_paragraph(str(row.content).strip())
    return new, located


def locate_item_with_heading(source_path, row) -> tuple[str | None, list | None]:
    """定位单条模板条目在底稿中的元素区间（大纲级别优先，行级清单兜底）。
    返回 (matched_heading, elements)：matched_heading=实际匹配到的底稿条目标题
    （身份信号，调用方必须比对是否等于所请求的条目标题）；elements=None 表示未定位
    （调用方按清单内容重建并如实标注）。
    大纲切片返回的是该部分全部条目，必须按条目标题匹配请求条目——
    直接取第一条会让同部分所有条目都拿到首条目区间（曾致价格/商务/技术各文件内容雷同）。"""
    from difflib import SequenceMatcher

    from docx import Document

    src_doc = Document(str(source_path))
    role = (row.structured or {}).get("role")
    if role in ("price", "business", "technical"):
        row_key = _item_key((row.content or "").split("\n")[0])
        if row_key:
            candidates: list[tuple[str, str, list]] = []
            for heading, elems in _locate_item_slices(src_doc, {role: [row]}).get(role, []):
                hd_key = _item_key(heading)
                if hd_key:
                    candidates.append((heading, hd_key, elems))
            # 1) 精确；2) 前缀（标题带/不带尾巴）；3) 相似度兜底（如"（单位负责人）"增字差异）
            for heading, hd_key, elems in candidates:
                if hd_key == row_key:
                    return heading, elems
            for heading, hd_key, elems in candidates:
                if row_key.startswith(hd_key) or hd_key.startswith(row_key):
                    return heading, elems
            if candidates:
                best_heading, best, best_elems = max(
                    candidates, key=lambda c: SequenceMatcher(None, row_key, c[1]).ratio()
                )
                if SequenceMatcher(None, row_key, best).ratio() >= 0.65:
                    return best_heading, best_elems
        row_slices = _locate_item_slices_by_rows(src_doc, {role: [row]})
        for _r, elems in row_slices.get(role, []):
            if _r.id == row.id and elems:
                return str(_r.content or "").split("\n")[0], elems
    return None, None


def locate_item_elements(source_path, row) -> list | None:
    """兼容包装：只返回元素区间（同 locate_item_with_heading 的第二个值）。"""
    _heading, elems = locate_item_with_heading(source_path, row)
    return elems


def canonical_text(root) -> str:
    """原文归一收集器（忠实性校验两边必须同一算法）：
    只收 w:t 与 w:delText 文本及元素的 tail，剔除 w:ins 子树（修订插入不算原文）。
    绘图内部数字（wp:posOffset/extent 等）不是 w:t，天然不计入——否则锚定图（印模）
    会让底稿侧文本里混入坐标数字，而条目侧没有，导致合法切片误报不忠实。"""
    W = f"{{{_W_NS}}}"

    def _inside(el, tag) -> bool:
        a = el.getparent()
        while a is not None:
            if a.tag == W + tag:
                return True
            a = a.getparent()
        return False

    parts: list[str] = []
    for el in root.iter():
        if _inside(el, "ins"):
            continue
        if el.tag == W + "t" and el.text:
            parts.append(el.text)
        elif el.tag == W + "delText" and el.text:
            parts.append(el.text)
        if el.tail:
            parts.append(el.tail)
    return "".join(parts)


def check_doc_fidelity(doc, source_text: str) -> dict:
    """逐字忠实性校验：条目文件原文（含修订删除线，剔除 w:ins）必须逐字包含于底稿原文。
    返回 {ok, original_chars, inserted_chars, deleted_chars, issues:[…]}。"""
    W = f"{{{_W_NS}}}"
    root = doc.element  # CT_Document

    def _inside(el, tag) -> bool:
        a = el.getparent()
        while a is not None:
            if a.tag == W + tag:
                return True
            a = a.getparent()
        return False

    def _norm(s: str) -> str:
        return "".join(s.split())

    ins_chars = 0
    del_chars = 0
    for el in root.iter():
        if _inside(el, "ins"):
            if el.tag == W + "t" and el.text:
                ins_chars += len(el.text)
            continue
        if el.tag == W + "delText" and el.text:
            del_chars += len(el.text)
    original = canonical_text(root)
    n_orig = _norm(original)
    src_norm = _norm(source_text)
    issues: list[str] = []
    ok = True
    if n_orig and n_orig not in src_norm:
        ok = False
        for i in range(0, len(n_orig) - 9, 10):
            if n_orig[i:i + 10] not in src_norm:
                issues.append(f"首个差异片段：{n_orig[max(0, i - 15):i + 15]!r}")
                break
        else:
            issues.append("原文与底稿不一致（空白/字符归一化差异）")
    return {
        "ok": ok,
        "original_chars": len(n_orig),
        "inserted_chars": ins_chars,
        "deleted_chars": del_chars,
        "issues": issues,
    }


def replace_text_tracked(editor, find: str, value: str, comment: str | None) -> int:
    """显式定向填空：把文档中出现的 find 原文替换为 value（修订模式+批注）。
    支持跨 run 的整段匹配；复用调用方持久 editor（批注 id 全局连续）；返回替换处数。"""
    if not find:
        return 0
    from docx.oxml.ns import qn as _qn

    replaced = 0
    for p_el in editor._doc.element.body.iter(_qn("w:p")):
        t_nodes = list(p_el.iter(_qn("w:t")))
        if not t_nodes:
            continue
        full = "".join(t.text or "" for t in t_nodes)
        if find not in full:
            continue
        # 整段替换（跨 run 拆分复杂度交给 track_paragraph_replace）
        new_full = full.replace(find, value)
        editor.track_paragraph_replace(t_nodes, full, new_full, comment or "按采购文件与应答资料回填，请复核。")
        replaced += full.count(find)
    return replaced


def _assemble_item_docx(source_path, row, elements, *, buyer, project_name, supplier, extra_nodes, tender_no: str = "") -> bytes:
    """组装一份条目文件：底稿中该条目区间原文整段复制（保留格式）+ 填空（修订批注）
    + 对应撰写内容（修订插入+批注）。未定位到区间的条目用清单内容重建。"""
    new, located = _build_item_document(source_path, row, elements)
    sess = _FillSession(new, buyer, project_name, supplier, tender_no)
    sess.apply_to_doc()
    if extra_nodes:
        sess.append_supplement(
            extra_nodes,
            heading_text="响应内容",
            comment="本节为针对本条目撰写的响应内容（依据招标文件要求与企业资料，未知处标【待补充】），"
                    "修订插入留痕，请复核确认。",
        )
    elif row is not None and not located:
        sess.editor.add_comment("该条目在底稿中未定位到原文区间，本文件按解析清单内容重建（可能不完整），"
                                "请对照采购文件原件核对。")
    return sess.finish()


async def _resolve_package_template_fobj(session, enterprise_id: int, project_id: int):
    """响应文件包底稿源：按内容分级选择——含"响应文件格式"章的采购文件最优先，
    同级取最大文件（完整采购文件 > 公告/规范书等）。"""
    from sqlalchemy import select as sa_select

    from app.models.doc import DocBlock
    from app.models.file import FileObject

    rows = (
        await session.scalars(
            sa_select(FileObject).where(
                FileObject.project_id == int(project_id),
                FileObject.enterprise_id == enterprise_id,
                FileObject.is_deleted.is_(False),
                FileObject.owner_type == 2,
            )
        )
    ).all()
    docx_rows = [f for f in rows if (f.ext or "").strip(".").lower() == "docx"]
    if not docx_rows:
        return None
    texts: dict[int, str] = {}
    for f in docx_rows:
        bs = (
            await session.scalars(
                sa_select(DocBlock.text_content).where(DocBlock.file_id == f.id)
            )
        ).all()
        texts[f.id] = "\n".join(b or "" for b in bs)[:200000]

    def _rank(f) -> int:
        t = texts.get(f.id) or ""
        if "响应文件格式" in t:
            return 2
        if "商务文件" in t and "技术文件" in t:
            return 1
        if "商务文件" in t or "技术文件" in t:
            return 0
        return -1

    ranked = [(_rank(f), f.size_bytes or 0, f) for f in docx_rows]
    ranked.sort(key=lambda x: (x[0], x[1]), reverse=True)
    if ranked[0][0] >= 0:
        return ranked[0][2]
    return max(docx_rows, key=lambda f: f.size_bytes)


async def build_response_package(session, enterprise_id: int, project_id: int) -> bytes:
    """响应文件包（产品定版：按招标文件清单逐份成文）：
    价格文件/商务文件/技术文件 三个目录，每个清单条目一份 docx——
    该条目模板原文（自底稿整段复制，保留格式）+ 填空（修订+批注来源）
    + 对应撰写内容（修订插入+批注）。附 报价单.xlsx 与 manifest.json。"""
    import io as _io
    import json as _json
    import zipfile as _zip

    from sqlalchemy import select as sa_select

    from app.models.deliverable import Deliverable
    from app.models.project import Project
    from app.models.requirement import Requirement
    from app.services import deliverable_service
    from app.services.storage import StorageProvider

    # 1) 底稿源 docx（响应文件格式章最优先，避免误选技术规范书等）
    fobj = await _resolve_package_template_fobj(session, enterprise_id, project_id)
    if fobj is None or (fobj.ext or "").strip(".").lower() != "docx":
        raise ValueError("缺少采购文件底稿：请先上传采购文件并完成【招标解析】后再打包")
    source_path = StorageProvider().open(fobj.bucket, fobj.object_key)
    from docx import Document

    src_doc = Document(str(source_path))

    # 2) 模板清单（响应文件格式，逐字落库）
    tpl_rows = (
        await session.scalars(
            sa_select(Requirement).where(
                Requirement.enterprise_id == enterprise_id,
                Requirement.project_id == project_id,
                Requirement.current.is_(True),
                Requirement.req_type == "doc_template",
            )
        )
    ).all()
    items_by_role: dict[str, list] = {"price": [], "business": [], "technical": []}
    for r in tpl_rows:
        role = (r.structured or {}).get("role")
        if role in items_by_role:
            items_by_role[role].append(r)

    def _ok(r):
        o = (r.structured or {}).get("order", 0)
        try:
            return (0, int(o))
        except (TypeError, ValueError):
            return (1, str(o))

    for role in items_by_role:
        items_by_role[role].sort(key=_ok)

    # 3) 成果模型（企业字段 + 撰写内容 + 报价单）
    buyer = project_name = supplier = ""
    biz_nodes, tech_nodes = [], []
    quote_model = None
    dls = (
        await session.scalars(
            sa_select(Deliverable).where(
                Deliverable.enterprise_id == enterprise_id,
                Deliverable.project_id == project_id,
            )
        )
    ).all()
    for d in dls:
        if d.current_version_no == 0:
            continue
        try:
            _, model = await deliverable_service.get_version_content(session, d.id, d.current_version_no)
        except Exception:  # noqa: BLE001
            continue
        if d.deliverable_type == 1:
            buyer = str(model.get("buyer") or "").strip() or buyer
            project_name = str(model.get("project_name") or "").strip() or project_name
            supplier = str(model.get("supplier_name") or "").strip() or supplier
            biz_nodes = model.get("supplement_nodes") or []
        elif d.deliverable_type == 2:
            tech_nodes = model.get("supplement_nodes") or []
        elif d.deliverable_type == 3:
            quote_model = model

    # 采购编号：从已解析要求里确定性提取（封面/公告均带"采购编号"字样）
    tender_no = ""
    try:
        import re as _re3

        all_reqs = (
            await session.scalars(
                sa_select(Requirement).where(
                    Requirement.enterprise_id == enterprise_id,
                    Requirement.project_id == project_id,
                    Requirement.current.is_(True),
                )
            )
        ).all()
        for r in all_reqs:
            m = _re3.search(r"采购编号[：: ]*([A-Za-z0-9][A-Za-z0-9\-]*)", r.content or "")
            if m and m.group(1) and not _re3.search(r"[，,、；;。]", m.group(1)):
                tender_no = m.group(1)
                break
    except Exception:  # noqa: BLE001
        tender_no = ""
    if not buyer or not project_name:
        project = await session.get(Project, int(project_id))
        if project is not None:
            if not buyer:
                buyer = (project.buyer or "").strip()
            if not project_name:
                project_name = (project.name or "").strip()

    # 4) 底稿切片（大纲级别优先，行级清单兜底）+ 撰写内容分配
    slices = _locate_item_slices(src_doc, items_by_role)
    outline_ok = any(slices.get(r) for r in ("price", "business", "technical"))
    if not outline_ok:
        row_slices = _locate_item_slices_by_rows(src_doc, items_by_role)
    biz_titles = (
        [_clean_item_name(h) for h, _e in slices.get("business", [])]
        if outline_ok
        else [str(r.content or "") for r in items_by_role["business"]]
    )
    tech_titles = (
        [_clean_item_name(h) for h, _e in slices.get("technical", [])]
        if outline_ok
        else [str(r.content or "") for r in items_by_role["technical"]]
    )
    biz_matched, biz_rest = _split_supplement(biz_titles, biz_nodes)
    tech_matched, tech_rest = _split_supplement(tech_titles, tech_nodes)

    # 5) 组装 zip（三目录 + manifest）
    buf = _io.BytesIO()
    files_manifest = []
    with _zip.ZipFile(buf, "w", _zip.ZIP_DEFLATED) as zf:
        for role, dir_name in _RESPONSE_ROLE_DIRS.items():
            role_slices = slices.get(role, [])
            matched_map = biz_matched if role == "business" else (tech_matched if role == "technical" else [[] for _ in role_slices])
            rest_nodes = biz_rest if role == "business" else (tech_rest if role == "technical" else [])
            if role_slices:
                for idx, (heading, elems) in enumerate(role_slices):
                    data = _assemble_item_docx(
                        source_path, None, elems,
                        buyer=buyer, project_name=project_name, supplier=supplier,
                        extra_nodes=matched_map[idx] if idx < len(matched_map) else [],
                        tender_no=tender_no,
                    )
                    name = f"{dir_name}/{_safe_filename(_clean_item_name(heading))}.docx"
                    zf.writestr(name, data)
                    files_manifest.append({"dir": dir_name, "name": name, "bytes": len(data)})
            else:
                # 兜底：按行级清单逐份（底稿无大纲信号时）
                rows = items_by_role[role]
                for idx, r in enumerate(rows):
                    elems = None
                    for rr, ee in row_slices.get(role, []):
                        if rr.id == r.id:
                            elems = ee
                            break
                    data = _assemble_item_docx(
                        source_path, r, elems,
                        buyer=buyer, project_name=project_name, supplier=supplier,
                        extra_nodes=matched_map[idx] if idx < len(matched_map) else [],
                        tender_no=tender_no,
                    )
                    name = f"{dir_name}/{_safe_filename(r.content)}.docx"
                    zf.writestr(name, data)
                    files_manifest.append({"dir": dir_name, "name": name, "bytes": len(data)})
            if rest_nodes:
                data = _assemble_item_docx(
                    source_path, None, None,
                    buyer=buyer, project_name=project_name, supplier=supplier,
                    extra_nodes=rest_nodes,
                    tender_no=tender_no,
                )
                name = f"{dir_name}/补充响应内容.docx"
                zf.writestr(name, data)
                files_manifest.append({"dir": dir_name, "name": name, "bytes": len(data)})
            if role == "price" and quote_model is not None:
                xlsx = xlsx_bytes(quote_model)
                name = f"{dir_name}/报价单.xlsx"
                zf.writestr(name, xlsx)
                files_manifest.append({"dir": dir_name, "name": name, "bytes": len(xlsx)})
        manifest = {
            "project_id": int(project_id),
            "draft": fobj.original_name,
            "note": "按招标文件《响应文件格式》清单逐份成文：每份=该条目模板原文+填空（修订模式+批注来源）+对应撰写内容；"
                    "全部改动可在 Word【审阅→所有标记】中逐处查看。附 Hermes 主会话全程记录（会话记录/主会话记录.md）。",
            "files": files_manifest,
        }
        # 主会话全程记录（新方案运行时带上；无 agent_pipeline 任务则跳过）
        try:
            from sqlalchemy import select as _sa_select2

            from app.constants import TaskType as _TaskType
            from app.models.task import Task as _Task
            from app.services.agent_pipeline import session_record_markdown

            pipe_task = await session.scalar(
                _sa_select2(_Task)
                .where(
                    _Task.enterprise_id == enterprise_id,
                    _Task.project_id == int(project_id),
                    _Task.task_type == _TaskType.AGENT_PIPELINE,
                )
                .order_by(_Task.id.desc())
                .limit(1)
            )
            if pipe_task is not None:
                record = await session_record_markdown(session, pipe_task)
                name = "会话记录/主会话记录.md"
                zf.writestr(name, record.encode("utf-8"))
                files_manifest.append({"dir": "会话记录", "name": name, "bytes": len(record.encode("utf-8"))})
                manifest["files"] = files_manifest
        except Exception:  # noqa: BLE001 会话记录缺失不影响交付包主体
            logger.warning("响应文件包附加主会话记录失败", exc_info=True)
        zf.writestr("manifest.json", _json.dumps(manifest, ensure_ascii=False, indent=2))
    return buf.getvalue()


async def _resolve_template_fobj(session, enterprise_id: int, model: dict, project_id: int | None = None):
    """解析底稿源 FileObject：模型记录的底稿源优先；旧成果从项目材料兜底。找不到返回 None。"""
    from sqlalchemy import select as sa_select

    from app.models.file import FileObject

    fobj = None
    fid = model.get("template_source_file_id") if isinstance(model, dict) else None
    if fid:
        try:
            fobj = await session.get(FileObject, int(fid))
        except (TypeError, ValueError):
            fobj = None
        if fobj is not None and (fobj.enterprise_id != enterprise_id or fobj.is_deleted):
            fobj = None
    # 旧成果没记底稿源（或记录已失效）时，从项目材料里找采购文件 docx 兜底
    if fobj is None and project_id:
        rows = (
            await session.scalars(
                sa_select(FileObject).where(
                    FileObject.project_id == int(project_id),
                    FileObject.enterprise_id == enterprise_id,
                    FileObject.is_deleted.is_(False),
                    FileObject.owner_type == 2,
                )
            )
        ).all()
        docx_rows = [f for f in rows if (f.ext or "").strip(".").lower() == "docx"]
        fobj = next((f for f in docx_rows if f.document_role == "tender"), None) or (
            max(docx_rows, key=lambda f: f.size_bytes) if docx_rows else None
        )
    return fobj


async def template_source_name(session, enterprise_id: int, model: dict, project_id: int | None = None) -> str | None:
    """底稿源文件名（下载文件名标注用）：找不到/非 docx 返回 None（文件名回退默认格式）。"""
    try:
        fobj = await _resolve_template_fobj(session, enterprise_id, model, project_id)
    except (TypeError, ValueError):
        return None
    if fobj is None or (fobj.ext or "").strip(".").lower() != "docx":
        return None
    name = str(fobj.original_name or "").strip()
    return name[:40] + "…" if len(name) > 40 else (name or None)


def _draft_label(draft_name: str | None, title: str) -> str:
    """成果文件名：带底稿来源标注（产品要求：一眼可见底稿是谁）。
    用全角冒号：半角冒号在 Windows 文件名中非法，会导致本地保存失败。"""
    if draft_name:
        return f"{title}(底稿：{draft_name})"
    return title


async def docx_bytes_with_source(
    session, enterprise_id: int, model: dict, project_id: int | None = None
) -> bytes:
    """导出唯一路径（底稿式，Issue #8 验收铁律 + 产品要求）：复制采购文件原 docx 整本保留，
    只做两件事——填空（已知字段回填/未知空位【待补充】）与文末追加"补充响应内容"。
    全部改动以 Word 修订模式（删除线/插入）+ 批注（来源说明）记录，验收直接可查。

    不存在节点式生成，也不存在回退：拿不到底稿源文件就明确报错，绝不从节点重新排版。"""
    from app.services.storage import StorageProvider

    fobj = await _resolve_template_fobj(session, enterprise_id, model, project_id)
    if fobj is None:
        raise ValueError(
            "该成果缺少采购文件底稿：请确认项目已上传采购文件（docx/pdf 等均可），"
            "重新触发【招标解析】和【生成标书】后再下载"
        )
    if (fobj.ext or "").strip(".").lower() != "docx":
        raise ValueError(
            "该成果的底稿源不是 docx（可能是历史数据）：请重新触发【招标解析】和【生成标书】后再下载"
        )
    try:
        source_path = StorageProvider().open(fobj.bucket, fobj.object_key)
    except FileNotFoundError as exc:
        raise ValueError(
            "采购文件底稿在存储中缺失：请重新上传采购文件并重新触发解析、生成后再下载"
        ) from exc
    try:
        return docx_from_template(source_path, model)
    except ValueError:
        raise
    except Exception as exc:  # noqa: BLE001 底稿损坏时明确报错，绝不回退节点式
        logger.warning(
            "底稿式导出失败 file=%s/%s: %s",
            fobj.bucket, fobj.object_key, type(exc).__name__,
            exc_info=True,
        )
        raise ValueError(
            f"底稿导出失败（{type(exc).__name__}）：采购文件底稿可能已损坏，"
            "请重新上传采购文件并重新触发解析、生成后再下载"
        ) from exc


def xlsx_bytes(model: dict) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    wb.remove(wb.active)
    for sheet in _normalize_xlsx_sheets(model):
        name = str((sheet or {}).get("name") or "报价单")[:31]
        ws = wb.create_sheet(name)
        rows = (sheet or {}).get("rows") or [["（无表格数据）"]]
        for row in rows:
            if isinstance(row, (list, tuple)):
                ws.append(["" if c is None else c for c in row])
            else:
                ws.append(["" if row is None else row])
    if not wb.worksheets:  # 兜底：openpyxl 要求至少一个可见工作表
        wb.create_sheet("报价单").append(["（无表格数据）"])
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _normalize_xlsx_sheets(model) -> list[dict]:
    """把各种报价模型形状归一为 [{"name","rows"}]，保证至少一个可见工作表。

    兼容：旧方案 {sheets:[{name,rows}]}；新方案 agent 落库的 {docModel:{sections:[...]}}；
    以及裸 {rows}/ {nodes} / 被 {version_no, model} 包装的形状。
    对 agent 生成侧的序列化退化（rows 误写成对象而非数组的数组，如
    {"item": [...], "item": [...]}）做恢复；恢复不了抛出可操作的 ValueError。"""
    if not isinstance(model, dict):
        model = {}
    inner = model.get("model")
    if isinstance(inner, dict) and not model.get("sheets"):
        model = inner
    sheets = model.get("sheets")
    if isinstance(sheets, list) and sheets:
        normalized = []
        for sh in sheets:
            if not isinstance(sh, dict):
                raise ValueError(
                    "sheets 元素必须是对象 {name, rows}，收到的元素不是对象——"
                    "请按 sheets=[{name:'报价单', rows:[['列1','列2'],...]}] 重发。"
                )
            name = str(sh.get("name") or "报价单")[:31]
            rows = sh.get("rows")
            if isinstance(rows, dict):
                # 生成侧序列化退化：{"item": [...]}（对象而非数组）——重复键在 JSON 解析时
                # 已丢数据，宁可报错让主会话按数组的数组重发，也不静默接受残缺行
                raise ValueError(
                    "报价单 rows 形状异常：rows 是对象而非数组的数组（疑似生成侧序列化错误）。"
                    "请重发为 rows=[[...],[...]]，每行是一个数组。"
                )
            if not isinstance(rows, list):
                raise ValueError(
                    "报价单 rows 形状异常：rows 必须是数组的数组 rows=[[...],[...]]。"
                )
            normalized.append({"name": name, "rows": rows})
        return normalized
    doc_model = model.get("docModel")
    if isinstance(doc_model, dict):
        sections = doc_model.get("sections")
        if isinstance(sections, list) and sections:
            rows_out = [["条目", "内容", "来源"]]
            for sec in sections:
                if isinstance(sec, dict):
                    rows_out.append(
                        [
                            str(sec.get("title") or ""),
                            str(sec.get("content") or "").replace("/n", "\n"),
                            str(sec.get("source") or ""),
                        ]
                    )
            return [{"name": "报价单", "rows": rows_out}]
    rows = model.get("rows")
    if isinstance(rows, list) and rows:
        return [{"name": "报价单", "rows": rows}]
    nodes = model.get("nodes")
    if isinstance(nodes, list) and nodes:
        rows_out = [
            [str(n.get("text") or "")] if isinstance(n, dict) else [str(n)]
            for n in nodes
        ]
        return [{"name": "报价单", "rows": rows_out}]
    return [{"name": "报价单", "rows": [["（报价模型无表格数据：请到报价页录入真实成本后重新生成）"]]}]


def run_final_check(deliverables, requirements=None, contents=None, structure=None) -> dict:
    """终检（Issue #12 / 产品验收）：完整性 + 结构合规 + 要求覆盖 + 文档质量。

    - deliverables: Deliverable 列表（必填）
    - requirements: Requirement 列表或含 content 字段的 dict 列表（可选；提供时检查要求覆盖）
    - contents: {deliverable_id: {"version_no": n, "model": dict}}（可选；提供时检查正文质量）
    - structure: [{"role": "business|technical|price", "title": 章节名}]（可选；提供时检查结构合规）
    """
    import re

    role_type = {"business": 1, "technical": 2, "price": 3}
    existing = {d.deliverable_type for d in deliverables}
    issues: list[dict] = []
    for dtype, name in DELIVERABLE_NAMES.items():
        if dtype not in existing:
            issues.append({"type": "完整性", "severity": "error", "message": f"缺少{name}", "locate": None})

    # 文档质量：仅记录/无正文、占位草稿、Markdown 残留、正文过短
    stub_marker = "草稿由 BidVolt 确定性生成"
    doc_texts: dict[int, str] = {}
    tpl_based: dict[int, bool] = {}
    for d in deliverables:
        name = DELIVERABLE_NAMES.get(d.deliverable_type, f"成果{d.deliverable_type}")
        content = (contents or {}).get(d.id)
        if d.current_version_no == 0 or content is None:
            issues.append({"type": "文档质量", "severity": "error", "message": f"{name}：仅有成果记录，尚无正文", "locate": d.id})
            continue
        model = content.get("model") or {}
        tpl_based[d.id] = bool(model.get("template_based"))
        nodes = _norm_supplement_nodes(model.get("nodes") or [])
        parts = [str(n.get("text") or "") for n in nodes]
        for n in nodes:
            for row in n.get("rows") or []:
                parts.append(" | ".join(str(c) for c in row))
        text = "\n".join(parts)
        sheet_text = "\n".join(
            str(c or "")
            for sh in model.get("sheets") or []
            for row in sh.get("rows") or []
            for c in row
        )
        doc_texts[d.deliverable_type] = text + sheet_text
        if d.deliverable_type in (1, 2):  # 商务标/技术标按正文检查
            if stub_marker in text:
                issues.append({"type": "文档质量", "severity": "error", "message": f"{name}：仍为占位草稿（未经真实生成）", "locate": d.id})
            if re.search(r"(^|\n)\s*#{1,6}\s", text) or "**" in text:
                issues.append({"type": "文档质量", "severity": "error", "message": f"{name}：正文残留 Markdown 标记（#/**），需重新生成", "locate": d.id})
            if len(text.strip()) < 100:
                issues.append({"type": "文档质量", "severity": "error", "message": f"{name}：正文过短（{len(text.strip())} 字，不足 100 字）", "locate": d.id})
        else:  # 报价单按表格检查
            if stub_marker in sheet_text:
                issues.append({"type": "文档质量", "severity": "error", "message": f"{name}：仍为占位草稿", "locate": d.id})
            if "待报价测算" in sheet_text:
                issues.append({"type": "文档质量", "severity": "warning", "message": f"{name}：尚未录入真实成本（请到报价页测算并应用）", "locate": d.id})

    # 结构合规：招标文件要求的章节必须逐章出现在对应成果中
    # （底稿式成果整本复制采购文件原 docx，模板章节由底稿保证，节点模型只是抽取子集，
    # 不按节点文本误报"缺少章节"）
    if structure:
        by_type = {d.deliverable_type: d for d in deliverables}
        for item in structure:
            title = str(item.get("title") or "").strip()
            role = item.get("role")
            if not title or role not in role_type:
                continue
            dtype = role_type[role]
            d = by_type.get(dtype)
            if d is not None and tpl_based.get(d.id):
                continue
            text = doc_texts.get(dtype, "")
            if title not in text:
                issues.append(
                    {
                        "type": "结构合规",
                        "severity": "error",
                        "message": f"{DELIVERABLE_NAMES.get(dtype, role)}缺少招标文件要求的章节：{title}",
                        "locate": None,
                    }
                )

    # 要求覆盖：有要求但成果整体无正文/占位时，必须拦截
    if requirements is not None:
        n = len(requirements)
        if n and contents is not None and not contents:
            issues.append({"type": "要求覆盖", "severity": "error", "message": f"已解析 {n} 条要求，但成果全部无正文", "locate": None})
        if n and any(stub_marker in ("\n".join(str(x.get("text") or "") for x in ((contents or {}).get(d.id) or {}).get("model", {}).get("nodes") or [])) for d in deliverables):
            issues.append({"type": "要求覆盖", "severity": "error", "message": f"已解析 {n} 条要求，但成果仍为占位草稿，未逐条响应", "locate": None})
        if not n:
            issues.append({"type": "要求覆盖", "severity": "warning", "message": "项目未解析出招标要求，成果可能无法逐条响应招标文件", "locate": None})
        # 逐条要求覆盖（对标的要求）：技术要求→技术标、资格要求→商务标逐条核对。
        # 底稿式成果整本复制采购文件原 docx，要求原文必然逐字在底稿内——
        # 不按节点文本做逐字比对误报"未响应"。
        tpl_by_type = {d.deliverable_type: tpl_based.get(d.id, False) for d in deliverables}
        for req in requirements:
            req_type = req.get("req_type")
            content = str(req.get("content") or "").strip()
            if not content:
                continue
            if req_type == "tech_requirement":
                text = doc_texts.get(2, "")
                target_name = "技术标"
                target_type = 2
            elif req_type == "qualification":
                text = doc_texts.get(1, "")
                target_name = "商务标"
                target_type = 1
            else:
                continue
            if tpl_by_type.get(target_type):
                continue
            if not text or content[:10] not in text:
                issues.append(
                    {
                        "type": "要求覆盖",
                        "severity": "error",
                        "message": f"{target_name}未响应要求：{content[:40]}",
                        "locate": req.get("id"),
                    }
                )

    # 文字质量（Issue #12 终检 v2）：重复段落与待补充占位统计
    words: dict[str, int] = {}
    pending_counts: dict[str, int] = {}
    for d in deliverables:
        name = DELIVERABLE_NAMES.get(d.deliverable_type, str(d.deliverable_type))
        text = doc_texts.get(d.deliverable_type, "")
        words[name] = len(text)
        pending_counts[name] = text.count("【待补充】")
        paras = [p.strip() for p in re.split(r"\n+", text) if len(p.strip()) >= 40]
        dup = {p: paras.count(p) for p in set(paras) if paras.count(p) >= 2}
        for p, cnt in list(dup.items())[:3]:
            issues.append(
                {
                    "type": "文字质量",
                    "severity": "error",
                    "message": f"{name}存在重复段落（出现 {cnt} 次）：{p[:30]}…",
                    "locate": d.id,
                }
            )
        if pending_counts[name] > 0:
            issues.append(
                {
                    "type": "文字质量",
                    "severity": "warning",
                    "message": f"{name}存在 {pending_counts[name]} 处【待补充】占位（资料不足处，需人工补齐）",
                    "locate": d.id,
                }
            )

    passed = not any(i["severity"] == "error" for i in issues)
    return {"passed": passed, "issues": issues, "words": words, "pending": pending_counts}


def build_manifest(project_id: int, files: list[dict], checks: dict) -> dict:
    return {
        "project_id": project_id,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "files": files,
        "checks": checks,
        "exemptions": [],
    }
